#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
tenderctl: deterministic helpers for the tender writing workflow.

Supported operations:
- template-profile: inspect a DOCX template and export JSON + Markdown notes.
- ingest-inputs: extract tender, scoring-standard, and optional user-brief text.
- brief-check: compare user-provided solution brief against tender text.
- claim-section / complete-section / merge-sections: coordinate parallel section drafting.
- similarity: compare generated draft text with tender source text.
- draft-audit: check heading depth, scoring-title visibility, and page budget.
- build-docx: render tender Markdown into a DOCX template with explicit styles.
"""

from __future__ import annotations

import argparse
import json
import math
import re
import shutil
import sys
from copy import deepcopy
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable, Optional

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.shared import Inches
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except Exception:  # pragma: no cover - optional runtime dependency
    Document = None
    WD_STYLE_TYPE = None
    WD_ALIGN_PARAGRAPH = None
    OxmlElement = None
    Inches = None
    Table = object
    Paragraph = object

try:
    import fitz  # PyMuPDF
except Exception:  # pragma: no cover - optional runtime dependency
    fitz = None

try:
    import yaml
except Exception:  # pragma: no cover - optional runtime dependency
    yaml = None


DEFAULT_STYLE_MAP = {
    "h1": "Heading 1",
    "h2": "Heading 2",
    "h3": "Heading 3",
    "h4": "Heading 3",
    "paragraph": "Normal",
    "bullet": "List Bullet",
    "number": "List Number",
    "caption": "Caption",
    "code": "Intense Quote",
    "table": "Table Grid",
}

DEFAULT_CHARS_PER_PAGE_ESTIMATE = 1800


def require_docx() -> None:
    if Document is None:
        raise RuntimeError("python-docx is not installed. Install tools/tenderctl/requirements.txt first.")

WORKFLOW_STAGES = [
    ("init", "初始化工作区、模板和进度文件"),
    ("template", "注册/选择模板，导出模板 Markdown 和格式规则"),
    ("extract", "抽取招标要求、评分项、废标项和格式约束"),
    ("brief", "核对用户基础资料与招标文件的相关性和偏离风险"),
    ("outline", "先建矩阵，再生成提纲"),
    ("draft", "按章节增量写作，逐步推进进度"),
    ("review", "按 S0-S4 和 PASS/WARN/FAIL/MANUAL 审查"),
    ("build", "按模板样式生成 final.docx"),
]


def ensure_parent(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore") if path.exists() else ""


def write_text(path: Path, text: str) -> None:
    ensure_parent(path)
    path.write_text(text, encoding="utf-8")


def write_json(path: Path, obj: object) -> None:
    ensure_parent(path)
    path.write_text(json.dumps(obj, ensure_ascii=False, indent=2), encoding="utf-8")


def read_json(path: Path, default: object) -> object:
    if not path.exists():
        return default
    return json.loads(path.read_text(encoding="utf-8"))


def read_manifest(path: Path) -> dict:
    if not path.exists():
        return {}
    text = path.read_text(encoding="utf-8")
    if yaml is not None:
        return yaml.safe_load(text) or {}
    # Fallback for simple top-level manifests. Full nested YAML requires PyYAML.
    manifest: dict[str, object] = {}
    for line in text.splitlines():
        if not line.strip() or line.lstrip().startswith("#") or ":" not in line:
            continue
        key, value = line.split(":", 1)
        value = value.strip().strip('"').strip("'")
        if value.lower() in {"true", "false"}:
            parsed: object = value.lower() == "true"
        elif re.match(r"^-?\d+$", value):
            parsed = int(value)
        else:
            parsed = value
        manifest[key.strip()] = parsed
    return manifest


def write_manifest(path: Path, manifest: dict) -> None:
    ensure_parent(path)
    if yaml is not None:
        path.write_text(yaml.safe_dump(manifest, allow_unicode=True, sort_keys=False), encoding="utf-8")
    else:
        path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")


def normalize_ws(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def style_name(style) -> str:
    try:
        return style.name or ""
    except Exception:
        return ""


def pt_value(value) -> Optional[float]:
    if value is None:
        return None
    try:
        return round(value.pt, 2)
    except Exception:
        return None


def bool_value(value) -> Optional[bool]:
    return None if value is None else bool(value)


def float_value(value) -> Optional[float]:
    if value is None:
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def alignment_name(value) -> Optional[str]:
    if value is None:
        return None
    if WD_ALIGN_PARAGRAPH is None:
        return str(value)
    names = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    }
    return names.get(value, str(value))


def font_name(font) -> Optional[str]:
    if font is None:
        return None
    try:
        return font.name
    except Exception:
        return None


def style_snapshot(style) -> dict:
    font = getattr(style, "font", None)
    p_fmt = getattr(style, "paragraph_format", None)
    return {
        "name": style_name(style),
        "type": str(getattr(style, "type", "")),
        "base_style": style_name(getattr(style, "base_style", None)),
        "font": font_name(font),
        "size_pt": pt_value(getattr(font, "size", None)),
        "bold": bool_value(getattr(font, "bold", None)),
        "italic": bool_value(getattr(font, "italic", None)),
        "alignment": alignment_name(getattr(p_fmt, "alignment", None)) if p_fmt else None,
        "line_spacing": getattr(p_fmt, "line_spacing", None) if p_fmt else None,
        "space_before_pt": pt_value(getattr(p_fmt, "space_before", None)) if p_fmt else None,
        "space_after_pt": pt_value(getattr(p_fmt, "space_after", None)) if p_fmt else None,
        "first_line_indent_pt": pt_value(getattr(p_fmt, "first_line_indent", None)) if p_fmt else None,
        "left_indent_pt": pt_value(getattr(p_fmt, "left_indent", None)) if p_fmt else None,
    }


def estimate_line_height_pt(font_size_pt: float, line_spacing: object = None) -> float:
    spacing = float_value(line_spacing)
    if spacing is None:
        return font_size_pt * 1.35
    if spacing <= 0:
        return font_size_pt * 1.35
    if spacing < 10:
        return font_size_pt * spacing
    return spacing


def estimate_chars_per_page_from_layout(layout: dict) -> int:
    page_width = float_value(layout.get("page_width_pt")) or 595.3
    page_height = float_value(layout.get("page_height_pt")) or 841.9
    left_margin = float_value(layout.get("left_margin_pt")) or 72.0
    right_margin = float_value(layout.get("right_margin_pt")) or 72.0
    top_margin = float_value(layout.get("top_margin_pt")) or 72.0
    bottom_margin = float_value(layout.get("bottom_margin_pt")) or 72.0
    font_size = float_value(layout.get("font_size_pt")) or 10.5
    line_height = estimate_line_height_pt(font_size, layout.get("line_spacing"))
    space_before = max(float_value(layout.get("space_before_pt")) or 0.0, 0.0)
    space_after = max(float_value(layout.get("space_after_pt")) or 0.0, 0.0)

    content_width = max(page_width - left_margin - right_margin, font_size * 12)
    content_height = max(page_height - top_margin - bottom_margin, line_height * 12)
    chars_per_line = max(int(content_width / max(font_size, 1.0)), 10)
    effective_line_height = max(line_height + (space_before + space_after) * 0.12, font_size)
    lines_per_page = max(int(content_height / effective_line_height), 8)

    # Headings, paragraph spacing, tables, bullets, and figure captions consume space
    # that is not represented in plain effective character counts.
    efficiency = 0.84
    if line_height >= font_size * 1.45:
        efficiency = 0.80
    elif line_height <= font_size * 1.2:
        efficiency = 0.88

    estimate = int(round(chars_per_line * lines_per_page * efficiency))
    return max(700, min(2600, estimate))


def estimate_chars_per_page_from_template_profile(profile: dict) -> int:
    sections = profile.get("sections") or []
    section = sections[0] if sections and isinstance(sections[0], dict) else {}
    style_map = {}
    if isinstance(profile.get("format_rules"), dict):
        style_map = profile["format_rules"].get("style_map") or {}
    style_map.update(profile.get("style_map") or {})
    paragraph_style_name = style_map.get("paragraph") or "Normal"

    paragraph_style = {}
    for style in profile.get("styles") or []:
        if isinstance(style, dict) and style.get("name") == paragraph_style_name:
            paragraph_style = style
            break

    layout = dict(section)
    layout.update(
        {
            "font_size_pt": paragraph_style.get("size_pt") or 10.5,
            "line_spacing": paragraph_style.get("line_spacing"),
            "space_before_pt": paragraph_style.get("space_before_pt"),
            "space_after_pt": paragraph_style.get("space_after_pt"),
        }
    )
    return estimate_chars_per_page_from_layout(layout)


def paragraph_snapshot(paragraph: Paragraph, idx: int) -> dict:
    text = normalize_ws(paragraph.text)
    runs = []
    for run in paragraph.runs[:12]:
        if not run.text.strip():
            continue
        runs.append(
            {
                "text_preview": run.text.strip()[:80],
                "font": font_name(run.font),
                "size_pt": pt_value(run.font.size),
                "bold": bool_value(run.font.bold),
                "italic": bool_value(run.font.italic),
            }
        )
    return {
        "index": idx,
        "text_preview": text[:160],
        "style": style_name(paragraph.style),
        "alignment": alignment_name(paragraph.alignment),
        "runs": runs,
    }


def table_snapshot(table: Table, idx: int) -> dict:
    rows = []
    for row in table.rows[:4]:
        rows.append([normalize_ws(cell.text)[:80] for cell in row.cells[:8]])
    return {
        "index": idx,
        "style": style_name(table.style),
        "rows": len(table.rows),
        "cols": len(table.columns),
        "preview": rows,
    }


def detect_style_map(doc: Document) -> dict:
    names = {style_name(s) for s in doc.styles}
    style_map = deepcopy(DEFAULT_STYLE_MAP)

    candidates = {
        "h1": ["标题 1", "Heading 1", "1"],
        "h2": ["标题 2", "Heading 2", "2"],
        "h3": ["标题 3", "Heading 3", "3"],
        "paragraph": ["正文", "Normal", "Body Text"],
        "caption": ["题注", "Caption"],
        "bullet": ["列表项目符号", "List Bullet"],
        "number": ["列表编号", "List Number"],
    }
    for role, role_candidates in candidates.items():
        for candidate in role_candidates:
            if candidate in names:
                style_map[role] = candidate
                break
    return style_map


def extract_template_profile(template_path: Path, template_id: str) -> dict:
    require_docx()
    doc = Document(str(template_path))
    styles = []
    for style in doc.styles:
        try:
            if style.type in (WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.TABLE, WD_STYLE_TYPE.CHARACTER):
                styles.append(style_snapshot(style))
        except Exception:
            continue

    sections = []
    for idx, section in enumerate(doc.sections):
        sections.append(
            {
                "index": idx,
                "orientation": str(section.orientation),
                "page_width_pt": pt_value(section.page_width),
                "page_height_pt": pt_value(section.page_height),
                "top_margin_pt": pt_value(section.top_margin),
                "bottom_margin_pt": pt_value(section.bottom_margin),
                "left_margin_pt": pt_value(section.left_margin),
                "right_margin_pt": pt_value(section.right_margin),
                "header_distance_pt": pt_value(section.header_distance),
                "footer_distance_pt": pt_value(section.footer_distance),
            }
        )

    paragraphs = [paragraph_snapshot(p, i) for i, p in enumerate(doc.paragraphs)]
    tables = [table_snapshot(t, i) for i, t in enumerate(doc.tables)]
    placeholder = "{{TECH_SECTION_BODY}}"
    placeholder_indexes = [idx for idx, p in enumerate(doc.paragraphs) if placeholder in (p.text or "")]
    placeholder_found = bool(placeholder_indexes)
    style_map = detect_style_map(doc)

    profile = {
        "template_id": template_id,
        "template_docx": str(template_path),
        "template_source_docx": str(template_path),
        "template_build_docx": str(template_path),
        "result": "PASS" if placeholder_found else "FAIL",
        "body_placeholder": placeholder,
        "source_placeholder_found": placeholder_found,
        "placeholder_found": placeholder_found,
        "placeholder_paragraph_indexes": placeholder_indexes,
        "placeholder_insertion": {
            "mode": "source",
            "anchor": "",
            "paragraph_index": placeholder_indexes[0] if placeholder_indexes else None,
        },
        "issues": []
        if placeholder_found
        else [
            "Template DOCX does not contain a standalone {{TECH_SECTION_BODY}} placeholder; build-docx cannot safely insert generated tender content."
        ],
        "style_map": style_map,
        "sections": sections,
        "styles": styles,
        "paragraphs": paragraphs[:300],
        "tables": tables[:100],
        "format_rules": {
            "source": "template_docx",
            "style_map_source": "auto_detected_from_template",
            "style_map": style_map,
            "headings": "Markdown #/##/### are mapped to template heading styles. Do not use #### or deeper for scoring items.",
            "paragraphs": "Normal tender text is rendered with the template paragraph style.",
            "lists": "Markdown ordered/unordered lists are mapped to template list styles.",
            "tables": "Markdown pipe tables are rendered as Word tables with the configured table style.",
            "figures": "Images are inserted from work/50_figures/out or the markdown-relative path; captions use the caption style.",
            "manual_checks": [
                "Update the Word table of contents after build.",
                "Confirm page numbers in scoring indexes after final pagination.",
                "Confirm section breaks for landscape wide tables manually when needed.",
            ],
        },
    }
    chars_per_page_estimate = estimate_chars_per_page_from_template_profile(profile)
    profile["chars_per_page_estimate"] = chars_per_page_estimate
    profile["page_estimate"] = {
        "chars_per_page_estimate": chars_per_page_estimate,
        "method": "template_geometry",
        "note": "Estimated from page size, margins, paragraph font size, and line spacing. Final Word/PDF pagination still requires review.",
    }
    profile["format_rules"]["chars_per_page_estimate"] = chars_per_page_estimate
    profile["format_rules"]["page_estimate"] = profile["page_estimate"]
    return profile


def insert_paragraph_after(paragraph: Paragraph, text: str, style: Optional[str] = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    new_para.add_run(text)
    return new_para


DEFAULT_TEMPLATE_INSERT_ANCHORS = (
    "投标人项目实施和售后服务方案",
    "其他技术文件",
    "技术部分",
)

DEFAULT_TEMPLATE_BODY_START_ANCHORS = (
    "评分标准相关内容索引表",
    "应答索引汇总表",
    "投标人项目实施和售后服务方案",
    "其他技术文件",
    "技术部分",
)


def remove_paragraph_element(paragraph: Paragraph) -> None:
    parent = paragraph._element.getparent()
    if parent is not None:
        parent.remove(paragraph._element)


def remove_body_elements_between(start_paragraph: Paragraph, end_paragraph: Optional[Paragraph]) -> None:
    start_el = start_paragraph._element
    end_el = end_paragraph._element if end_paragraph is not None else None
    parent = start_el.getparent()
    current = start_el.getnext()
    while parent is not None and current is not None and current is not end_el:
        next_el = current.getnext()
        if current.tag.endswith("}sectPr"):
            break
        parent.remove(current)
        current = next_el


def clear_document_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def find_paragraph_index(
    paragraphs: list[Paragraph],
    anchors: list[str],
    start_at: int = 0,
    prefer_last: bool = True,
) -> tuple[Optional[int], str]:
    for anchor in anchors:
        if not anchor:
            continue
        matches = [
            idx
            for idx, paragraph in enumerate(paragraphs[start_at:], start=start_at)
            if anchor in normalize_ws(paragraph.text or "")
        ]
        if matches:
            return (matches[-1] if prefer_last else matches[0]), anchor
    return None, ""


def prepare_build_template(
    source_docx: Path,
    build_docx: Path,
    placeholder: str = "{{TECH_SECTION_BODY}}",
    insert_after: str = "",
    build_mode: str = "local_template",
    body_start_after: str = "",
    body_end_before: str = "",
) -> dict:
    require_docx()
    source_docx = Path(source_docx)
    build_docx = Path(build_docx)
    source_doc = Document(str(source_docx))
    source_placeholder_indexes = [
        idx for idx, paragraph in enumerate(source_doc.paragraphs) if placeholder in (paragraph.text or "")
    ]
    if source_placeholder_indexes:
        return {
            "template_source_docx": str(source_docx),
            "template_build_docx": str(source_docx),
            "source_placeholder_found": True,
            "placeholder_found": True,
            "placeholder_paragraph_indexes": source_placeholder_indexes,
            "placeholder_insertion": {
                "mode": "source",
                "anchor": "",
                "paragraph_index": source_placeholder_indexes[0],
            },
            "issues": [],
        }

    ensure_parent(build_docx)
    shutil.copyfile(source_docx, build_docx)
    build_doc = Document(str(build_docx))

    normalized_mode = (build_mode or "local_template").strip().lower()
    if normalized_mode not in {"local_template", "replace_body", "insert_after"}:
        normalized_mode = "local_template"

    issues = []
    anchor_used = ""
    end_used = ""
    mode = normalized_mode
    paragraph_index = None

    if normalized_mode == "local_template":
        clear_document_body(build_doc)
        build_doc.add_paragraph(placeholder)
        mode = "local_template"
        paragraph_index = 0
    elif normalized_mode == "replace_body":
        start_anchors = [body_start_after or insert_after] if (body_start_after or insert_after) else list(DEFAULT_TEMPLATE_BODY_START_ANCHORS)
        paragraphs = build_doc.paragraphs
        start_index, anchor_used = find_paragraph_index(paragraphs, start_anchors)
        if start_index is None:
            build_doc.add_paragraph(placeholder)
            mode = "append_end"
            issues.append(
                "No configured or default template body start anchor was found; {{TECH_SECTION_BODY}} was appended at the end of the build template. Set template_body_start_after or review template_build.docx before build."
            )
        else:
            end_index = len(paragraphs)
            if body_end_before:
                found_end, end_used = find_paragraph_index(
                    paragraphs, [body_end_before], start_at=start_index + 1, prefer_last=False
                )
                if found_end is not None:
                    end_index = found_end
                else:
                    issues.append(
                        f"template_body_end_before anchor was not found after {anchor_used}; body replacement continued to the end of the template."
                    )
            end_paragraph = paragraphs[end_index] if end_index < len(paragraphs) else None
            remove_body_elements_between(paragraphs[start_index], end_paragraph)
            insert_paragraph_after(paragraphs[start_index], placeholder)
            paragraph_index = start_index + 1
    else:
        anchors = [insert_after] if insert_after else list(DEFAULT_TEMPLATE_INSERT_ANCHORS)
        paragraphs = build_doc.paragraphs
        anchor_index, anchor_used = find_paragraph_index(paragraphs, anchors)
        if anchor_index is None:
            build_doc.add_paragraph(placeholder)
            mode = "append_end"
            issues.append(
                "No configured or default template insertion anchor was found; {{TECH_SECTION_BODY}} was appended at the end of the build template. Review template_build.docx before build."
            )
        else:
            insert_paragraph_after(paragraphs[anchor_index], placeholder)
            paragraph_index = anchor_index + 1

    build_doc.save(str(build_docx))
    build_doc = Document(str(build_docx))
    placeholder_indexes = [
        idx for idx, paragraph in enumerate(build_doc.paragraphs) if placeholder in (paragraph.text or "")
    ]
    return {
        "template_source_docx": str(source_docx),
        "template_build_docx": str(build_docx),
        "source_placeholder_found": False,
        "placeholder_found": bool(placeholder_indexes),
        "placeholder_paragraph_indexes": placeholder_indexes,
        "placeholder_insertion": {
            "mode": mode,
            "anchor": anchor_used,
            "start_after": anchor_used if normalized_mode == "replace_body" else "",
            "end_before": end_used,
            "paragraph_index": placeholder_indexes[0] if placeholder_indexes else None,
        },
        "issues": issues,
    }


def format_block(style: str, extra: Optional[dict] = None) -> str:
    payload = {"style": style}
    if extra:
        payload.update(extra)
    lines = ["<!-- tender:format"]
    for key, value in payload.items():
        lines.append(f"{key}: {value}")
    lines.append("-->")
    return "\n".join(lines)


def template_to_markdown(profile: dict) -> str:
    lines = [
        f"# 模板格式说明：{profile['template_id']}",
        "",
        "本文件由 Word 模板转换生成，用于说明模板结构和格式要求；不要把模板示例内容当作项目真实内容。",
        "",
        "## 样式映射",
        "",
        "| Markdown | Word 样式 | 用途 |",
        "|---|---|---|",
    ]
    style_map = profile.get("style_map", {})
    rows = [
        ("#", style_map.get("h1", "Heading 1"), "一级标题"),
        ("##", style_map.get("h2", "Heading 2"), "二级标题"),
        ("###", style_map.get("h3", "Heading 3"), "三级标题"),
        ("正文", style_map.get("paragraph", "Normal"), "正文段落"),
        ("- 列表", style_map.get("bullet", "List Bullet"), "无序列表"),
        ("1. 列表", style_map.get("number", "List Number"), "有序列表"),
        ("图/表题", style_map.get("caption", "Caption"), "题注"),
        ("表格", style_map.get("table", "Table Grid"), "表格"),
    ]
    for row in rows:
        lines.append("| " + " | ".join(row) + " |")

    lines.extend(["", "## 模板正文骨架", ""])
    if not profile.get("placeholder_found"):
        lines.extend(
            [
                "> [!warning] 模板缺少 `{{TECH_SECTION_BODY}}` 占位符，不能安全执行 `/tender:build`。请在 Word 模板中把该占位符单独放在正文插入位置，然后重新运行 `/tender:template`。",
                "",
            ]
        )
    elif not profile.get("source_placeholder_found"):
        insertion = profile.get("placeholder_insertion") or {}
        build_docx = profile.get("template_build_docx", "")
        anchor = insertion.get("anchor") or "模板末尾"
        if insertion.get("mode") == "local_template":
            anchor = "局部模板整体"
            mode_text = "作为写作边界，清空模板示例正文后插入占位符"
        elif insertion.get("mode") == "replace_body":
            mode_text = "替换正文骨架并插入占位符"
        else:
            mode_text = "插入占位符"
        lines.extend(
            [
                f"> [!info] 原始模板不含 `{{TECH_SECTION_BODY}}`；已生成构建模板 `{build_docx}`，并在 `{anchor}` 位置{mode_text}。`/tender:build` 默认使用构建模板，原始模板保持不变。",
                "",
            ]
        )
    if profile.get("issues"):
        lines.extend(["> [!warning] 需要复核：" + "；".join(str(item) for item in profile["issues"]), ""])
    for paragraph in profile.get("paragraphs", []):
        text = paragraph.get("text_preview", "").strip()
        style = paragraph.get("style") or "Normal"
        if not text:
            continue
        lines.append(format_block(style, {"source_paragraph": paragraph.get("index")}))
        if "{{" in text and "}}" in text:
            lines.append(text)
        elif style == style_map.get("h1"):
            lines.append(f"# {text}")
        elif style == style_map.get("h2"):
            lines.append(f"## {text}")
        elif style == style_map.get("h3"):
            lines.append(f"### {text}")
        else:
            lines.append(text)
        lines.append("")

    if profile.get("tables"):
        lines.extend(["## 表格样式样例", ""])
        for table in profile["tables"]:
            lines.append(format_block(table.get("style") or "Table Grid", {"source_table": table.get("index")}))
            preview = table.get("preview") or []
            if not preview:
                continue
            width = max(len(row) for row in preview)
            header = preview[0] + [""] * (width - len(preview[0]))
            lines.append("| " + " | ".join(header) + " |")
            lines.append("|" + "|".join(["---"] * width) + "|")
            for row in preview[1:]:
                row = row + [""] * (width - len(row))
                lines.append("| " + " | ".join(row) + " |")
            lines.append("")

    lines.extend(
        [
            "## 最终成文注意事项",
            "",
            "- 评分项应进入三级标题以内；不得为了套模板把评分项藏到四级标题以下。",
            "- 招标文件格式要求优先于公司模板；冲突时记录为 MANUAL 并请用户判断。",
            "- 生成 docx 后需要刷新目录，并复核评分索引表页码。",
        ]
    )
    return "\n".join(lines).rstrip() + "\n"


def cmd_template_profile(args) -> int:
    template_path = Path(args.template)
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    template_id = args.template_id or template_path.stem
    manifest_path = Path(args.manifest) if args.manifest else Path()
    existing_manifest = read_manifest(manifest_path) if args.manifest else {}
    out_json = Path(args.out_json)
    default_build_template = out_json.parent / "template_build.docx"
    build_template_path = Path(
        getattr(args, "build_template_out", "")
        or existing_manifest.get("template_build_docx_path")
        or default_build_template
    )
    insert_after = str(
        getattr(args, "insert_after", "") or existing_manifest.get("template_insert_after") or ""
    ).strip()
    arg_build_mode = str(getattr(args, "build_mode", "") or "").strip()
    manifest_build_mode = str(existing_manifest.get("template_build_mode") or "").strip()
    if arg_build_mode:
        build_mode = arg_build_mode
    elif manifest_build_mode:
        build_mode = manifest_build_mode
    elif insert_after:
        build_mode = "insert_after"
    else:
        build_mode = "local_template"
    template_scope = str(
        getattr(args, "template_scope", "")
        or existing_manifest.get("template_scope")
        or ("local" if build_mode == "local_template" else "complete")
    ).strip()
    body_start_after = str(
        getattr(args, "body_start_after", "")
        or existing_manifest.get("template_body_start_after")
        or (insert_after if build_mode != "local_template" else "")
        or ""
    ).strip()
    body_end_before = str(
        getattr(args, "body_end_before", "") or existing_manifest.get("template_body_end_before") or ""
    ).strip()
    prepared = prepare_build_template(
        source_docx=template_path,
        build_docx=build_template_path,
        placeholder="{{TECH_SECTION_BODY}}",
        insert_after=insert_after,
        build_mode=build_mode,
        body_start_after=body_start_after,
        body_end_before=body_end_before,
    )
    profile = extract_template_profile(template_path, template_id)
    profile.update(prepared)
    profile["template_scope"] = template_scope
    profile["template_docx"] = prepared["template_build_docx"]
    profile["result"] = "PASS" if prepared.get("placeholder_found") and not prepared.get("issues") else "MANUAL"
    if not prepared.get("placeholder_found"):
        profile["result"] = "FAIL"
    profile["issues"] = prepared.get("issues", [])

    out_md = Path(args.out_md)
    write_json(out_json, profile)
    write_text(out_md, template_to_markdown(profile))
    if getattr(args, "out_format_rules", ""):
        write_json(Path(args.out_format_rules), profile.get("format_rules", {}))
        print(f"[OK] wrote: {args.out_format_rules}")

    if args.manifest:
        manifest = read_manifest(manifest_path)
        templates = manifest.setdefault("templates", {})
        templates[template_id] = {
            "source_docx_path": str(template_path),
            "build_docx_path": str(prepared["template_build_docx"]),
            "docx_path": str(prepared["template_build_docx"]),
            "profile_json": str(out_json),
            "template_md": str(out_md),
            "template_scope": template_scope,
            "template_build_mode": build_mode,
        }
        manifest["active_template_id"] = template_id
        manifest["template_source_docx_path"] = str(template_path)
        manifest["template_build_docx_path"] = str(prepared["template_build_docx"])
        manifest["template_docx_path"] = str(prepared["template_build_docx"])
        manifest["template_scope"] = template_scope
        manifest["template_build_mode"] = build_mode or "local_template"
        manifest["template_body_start_after"] = body_start_after
        manifest["template_body_end_before"] = body_end_before
        manifest["template_profile_path"] = str(out_json)
        manifest["template_markdown_path"] = str(out_md)
        manifest["body_placeholder"] = prepared.get("body_placeholder", "{{TECH_SECTION_BODY}}")
        manifest["chars_per_page_estimate"] = int(profile.get("chars_per_page_estimate") or DEFAULT_CHARS_PER_PAGE_ESTIMATE)
        manifest["chars_per_page_source"] = "template_geometry"
        if insert_after:
            manifest["template_insert_after"] = insert_after
        write_manifest(manifest_path, manifest)

    print(f"[OK] wrote: {out_json}")
    print(f"[OK] wrote: {out_md}")
    return 0


def extract_text_from_pdf(pdf_path: Path) -> str:
    if fitz is None:
        raise RuntimeError("PyMuPDF is not installed. Install PyMuPDF or provide a DOCX tender file.")
    doc = fitz.open(str(pdf_path))
    parts = []
    for i in range(doc.page_count):
        parts.append(doc.load_page(i).get_text("text"))
    return normalize_ws("\n".join(parts))


def extract_text_from_docx(docx_path: Path) -> str:
    require_docx()
    doc = Document(str(docx_path))
    parts = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return normalize_ws("\n".join(parts))


def extract_text_from_any(input_path: Path) -> str:
    suffix = input_path.suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(input_path)
    if suffix == ".docx":
        return extract_text_from_docx(input_path)
    if suffix in {".md", ".txt"}:
        return normalize_ws(input_path.read_text(encoding="utf-8", errors="ignore"))
    raise ValueError(f"Unsupported input type: {input_path.suffix}")


def cmd_ingest(args) -> int:
    manifest = read_manifest(Path(args.manifest))
    tender_path = Path(args.tender or manifest.get("tender_input_path", "inputs/tender.pdf"))
    out_txt = Path(args.out_txt)

    if not tender_path.exists():
        raise FileNotFoundError(f"Tender input not found: {tender_path}")

    text = extract_text_from_any(tender_path)

    write_text(out_txt, text)
    print(f"[OK] wrote: {out_txt}")
    return 0


def cmd_ingest_inputs(args) -> int:
    manifest = read_manifest(Path(args.manifest))
    tender_path = Path(manifest.get("tender_input_path", "inputs/tender.pdf"))
    if not tender_path.exists():
        raise FileNotFoundError(f"Tender input not found: {tender_path}")

    tender_text = extract_text_from_any(tender_path)
    write_text(Path(args.out_tender_txt), tender_text)
    print(f"[OK] wrote: {args.out_tender_txt}")

    scoring_path_raw = str(manifest.get("scoring_standard_path", "") or "").strip()
    if scoring_path_raw:
        scoring_path = Path(scoring_path_raw)
        if not scoring_path.exists():
            raise FileNotFoundError(f"Scoring standard input not found: {scoring_path}")
        scoring_text = extract_text_from_any(scoring_path)
        write_text(Path(args.out_scoring_txt), scoring_text)
        print(f"[OK] wrote: {args.out_scoring_txt}")
    else:
        print("[OK] scoring_standard_path is empty; scoring standard should be extracted from tender.txt")

    user_brief_path_raw = str(manifest.get("user_brief_path", "") or "").strip()
    if user_brief_path_raw:
        user_brief_path = Path(user_brief_path_raw)
        if not user_brief_path.exists():
            if bool(manifest.get("user_brief_required", False)):
                raise FileNotFoundError(f"User brief input not found: {user_brief_path}")
            print(f"[OK] user_brief_path not found, skipped optional user brief: {user_brief_path}")
            return 0
        brief_text = extract_text_from_any(user_brief_path)
        write_text(Path(args.out_user_brief_txt), brief_text)
        print(f"[OK] wrote: {args.out_user_brief_txt}")
    else:
        print("[OK] user_brief_path is empty; no user brief extracted")
    return 0


def shingles(text: str, k: int = 10) -> set[str]:
    text = re.sub(r"\s+", " ", text).strip()
    if len(text) <= k:
        return {text} if text else set()
    return {text[i : i + k] for i in range(0, len(text) - k + 1)}


def jaccard(a: set[str], b: set[str]) -> float:
    if not a and not b:
        return 0.0
    return len(a & b) / len(a | b) if (a | b) else 0.0


def md_to_plain(markdown: str) -> str:
    markdown = re.sub(r"```.*?```", " ", markdown, flags=re.S)
    markdown = re.sub(r"<!--.*?-->", " ", markdown, flags=re.S)
    markdown = re.sub(r"!\[.*?\]\(.*?\)", " ", markdown)
    markdown = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", markdown)
    markdown = re.sub(r"[#>*`_\-|]+", " ", markdown)
    markdown = re.sub(r"\s+", " ", markdown)
    return markdown.strip()


def extract_terms(text: str) -> set[str]:
    text = re.sub(r"[^\w\u4e00-\u9fff]+", " ", text)
    raw_terms = re.findall(r"[A-Za-z][A-Za-z0-9_\-]{2,}|[\u4e00-\u9fff]{2,}", text)
    terms: set[str] = set()
    stop = {
        "本项目",
        "方案",
        "进行",
        "提供",
        "实现",
        "通过",
        "建设",
        "系统",
        "平台",
        "服务",
        "功能",
        "要求",
        "相关",
    }
    for term in raw_terms:
        if term in stop:
            continue
        if len(term) <= 8:
            terms.add(term.lower())
        else:
            for i in range(0, len(term) - 1):
                terms.add(term[i : i + 2].lower())
    return terms


def sentence_like_spans(text: str) -> list[str]:
    parts = re.split(r"[。！？!?；;\n]+", text)
    return [normalize_ws(p) for p in parts if len(normalize_ws(p)) >= 8]


def brief_risk_spans(brief_text: str, tender_text: str) -> list[dict]:
    tender_terms = extract_terms(tender_text)
    risk_words = [
        "必须采用",
        "统一采用",
        "全部采用",
        "源代码",
        "源码",
        "知识产权",
        "自主研发",
        "7x24",
        "7×24",
        "国产化",
        "等保",
        "性能",
        "并发",
        "响应时间",
        "到场",
        "免费",
        "承诺",
        "保证",
        "唯一",
    ]
    risks = []
    for span in sentence_like_spans(brief_text):
        if not any(word in span for word in risk_words):
            continue
        span_terms = extract_terms(span)
        overlap = len(span_terms & tender_terms)
        ratio = overlap / max(len(span_terms), 1)
        if ratio < 0.25:
            risks.append(
                {
                    "level": "MANUAL",
                    "reason": "用户资料包含承诺或技术路线表述，但与招标文件文本重合度较低，需要人工确认是否偏离。",
                    "text": span[:180],
                    "term_overlap_ratio": round(ratio, 4),
                }
            )
    return risks[:50]


def cmd_brief_check(args) -> int:
    tender_text = read_text(Path(args.tender_txt))
    brief_text = read_text(Path(args.user_brief_txt))
    if not tender_text:
        raise FileNotFoundError(f"Tender text is empty or missing: {args.tender_txt}")
    if not brief_text:
        raise FileNotFoundError(f"User brief text is empty or missing: {args.user_brief_txt}")

    tender_terms = extract_terms(tender_text)
    brief_terms = extract_terms(brief_text)
    shared = sorted(brief_terms & tender_terms)
    missing = sorted(brief_terms - tender_terms)
    relevance = len(shared) / max(len(brief_terms), 1)
    risks = brief_risk_spans(brief_text, tender_text)

    if relevance < float(args.warn_threshold):
        result = "WARN"
    else:
        result = "PASS"
    if risks:
        result = "MANUAL" if result == "PASS" else result

    report = {
        "result": result,
        "relevance_score": round(relevance, 4),
        "brief_terms": len(brief_terms),
        "shared_terms": shared[:120],
        "brief_terms_not_seen_in_tender": missing[:120],
        "risk_spans": risks,
        "policy": [
            "Tender requirements remain the highest-priority source.",
            "User brief can guide solution expansion only when aligned with tender requirements.",
            "Risk spans marked MANUAL must not be written as firm commitments before user confirmation.",
        ],
    }
    write_json(Path(args.out_json), report)

    md_lines = [
        "# 用户基础资料相关性检查",
        "",
        f"- 结果：{result}",
        f"- 相关性分数：{report['relevance_score']}",
        f"- 用户资料关键词数：{report['brief_terms']}",
        f"- 与招标文件重合关键词数：{len(shared)}",
        "",
        "## 可作为写作基础的重合关键词",
        "",
        "、".join(shared[:80]) if shared else "无明显重合关键词。",
        "",
        "## 可能偏离或需要确认的关键词",
        "",
        "、".join(missing[:80]) if missing else "暂无。",
        "",
        "## MANUAL 风险片段",
        "",
    ]
    if risks:
        for item in risks:
            md_lines.append(f"- {item['text']}（{item['reason']}）")
    else:
        md_lines.append("- 暂无明显高风险承诺片段。")
    md_lines.extend(
        [
            "",
            "## 使用规则",
            "",
            "- 后续写作应以招标文件为约束边界，以用户资料作为方案素材。",
            "- 用户资料中未被招标文件支撑的路线、指标、承诺，必须在正文中降级为可选方案或标记 MANUAL。",
            "- 用户中途调整资料后，应重新运行本检查并更新 outline/draft。",
        ]
    )
    write_text(Path(args.out_md), "\n".join(md_lines) + "\n")
    print(f"[OK] wrote: {args.out_json}")
    print(f"[OK] wrote: {args.out_md}")
    return 0


def cmd_similarity(args) -> int:
    tender_text = read_text(Path(args.tender_txt))
    draft_text = md_to_plain(read_text(Path(args.draft_md)))
    k = int(args.k)
    sim = jaccard(shingles(tender_text, k=k), shingles(draft_text, k=k))

    suspect = []
    window = int(args.window)
    step = int(args.step)
    threshold = float(args.threshold)
    for offset in range(0, max(len(draft_text) - window, 0), step):
        segment = draft_text[offset : offset + window]
        score = jaccard(shingles(tender_text, k=k), shingles(segment, k=k))
        if score >= threshold:
            suspect.append({"offset": offset, "score": round(score, 4), "preview": segment[:120]})

    report = {
        "overall_similarity": round(sim, 4),
        "threshold": threshold,
        "suspect_spans": suspect[:50],
        "note": "Similarity is only a local check against tender source text; it is not an internet plagiarism detector.",
    }
    write_json(Path(args.out_json), report)
    print(f"[OK] wrote: {args.out_json}")
    return 0


def strip_markdown_for_count(markdown: str) -> str:
    text = re.sub(r"```.*?```", " ", markdown, flags=re.S)
    text = re.sub(r"<!--.*?-->", " ", text, flags=re.S)
    text = re.sub(r"!\[[^\]]*\]\([^)]+\)", " ", text)
    text = re.sub(r"\[[^\]]+\]\([^)]+\)", " ", text)
    text = re.sub(r"^\s{0,3}#{1,6}\s+", "", text, flags=re.M)
    text = re.sub(r"[*_`>|#-]+", " ", text)
    return normalize_ws(text)


def effective_char_count(text: str) -> int:
    cjk = re.findall(r"[\u4e00-\u9fff]", text)
    words = re.findall(r"[A-Za-z0-9]+", text)
    other = re.sub(r"[\u4e00-\u9fffA-Za-z0-9\s]", "", text)
    return len(cjk) + sum(max(1, len(word) // 5) for word in words) + len(other)


def resolve_chars_per_page(value: object = None, manifest: Optional[dict] = None) -> int:
    for candidate in (value, (manifest or {}).get("chars_per_page_estimate"), DEFAULT_CHARS_PER_PAGE_ESTIMATE):
        try:
            parsed = int(candidate)
        except (TypeError, ValueError):
            continue
        if parsed > 0:
            return parsed
    return DEFAULT_CHARS_PER_PAGE_ESTIMATE


def resolve_chars_per_page_from_args(args) -> int:
    manifest_path = Path(getattr(args, "manifest", "work/00_manifest.yml") or "work/00_manifest.yml")
    manifest = read_manifest(manifest_path) if manifest_path.exists() else {}
    return resolve_chars_per_page(getattr(args, "chars_per_page", ""), manifest)


def normalize_title(text: str) -> str:
    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"\s+", "", text)
    text = re.sub(r"[，。、“”‘’：:；;（）()\[\]【】《》<>.,/\\|+\-*#_`~!！?？]", "", text)
    return text.lower()


def extract_md_headings(markdown: str) -> list[dict]:
    headings = []
    for line_no, line in enumerate(markdown.splitlines(), start=1):
        match = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if not match:
            continue
        headings.append(
            {
                "line": line_no,
                "level": len(match.group(1)),
                "title": match.group(2).strip(),
                "normalized": normalize_title(match.group(2)),
            }
        )
    return headings


def repeated_document_shell_headings(headings: list[dict]) -> list[dict]:
    shell_titles = {"投标文件": "投标文件", "目录": "目录"}
    occurrences: dict[str, list[dict]] = {}
    for heading in headings:
        if heading.get("level") != 1:
            continue
        normalized = str(heading.get("normalized") or "")
        if normalized not in shell_titles:
            continue
        occurrences.setdefault(normalized, []).append(heading)
    repeated = []
    for normalized, items in occurrences.items():
        if len(items) > 1:
            repeated.append(
                {
                    "title": shell_titles[normalized],
                    "count": len(items),
                    "lines": [item["line"] for item in items],
                }
            )
    return repeated


def parse_markdown_table(text: str) -> list[dict]:
    rows = []
    lines = [line.strip() for line in text.splitlines() if line.strip().startswith("|")]
    i = 0
    while i + 1 < len(lines):
        header = split_table_row(lines[i])
        if not is_table_separator(lines[i + 1]):
            i += 1
            continue
        i += 2
        while i < len(lines) and lines[i].startswith("|"):
            cells = split_table_row(lines[i])
            if len(cells) < len(header):
                cells.extend([""] * (len(header) - len(cells)))
            rows.append({header[j].strip(): cells[j].strip() for j in range(len(header))})
            i += 1
        break
    return rows


def extract_scoring_items(scoring_matrix_md: Path) -> list[str]:
    if not scoring_matrix_md.exists():
        return []
    rows = parse_markdown_table(read_text(scoring_matrix_md))
    items = []
    for row in rows:
        item = row.get("评分项") or row.get("评审项目") or row.get("评分标准号") or ""
        item = normalize_ws(item)
        if item and item not in items:
            items.append(item)
    return items


def number_value(value) -> Optional[float]:
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value).strip()
    if not text:
        return None
    try:
        return float(text)
    except ValueError:
        return None


def first_present(row: dict, candidates: Iterable[str]) -> str:
    normalized = {normalize_title(key): value for key, value in row.items()}
    for candidate in candidates:
        value = row.get(candidate)
        if value:
            return normalize_ws(str(value))
        value = normalized.get(normalize_title(candidate))
        if value:
            return normalize_ws(str(value))
    return ""


def row_contains_any(row: dict, needles: Iterable[str]) -> bool:
    text = normalize_title(" ".join(str(value) for value in row.values()))
    return any(normalize_title(needle) in text for needle in needles if needle)


def section_matches_text(section: dict, text: str) -> bool:
    if not text:
        return False
    normalized_text = normalize_title(text)
    section_id = section_id_of(section)
    title = str(section.get("title") or "")
    if section_id and normalize_title(section_id) in normalized_text:
        return True
    if title and normalize_title(title) in normalized_text:
        return True
    return False


def load_page_allocations(page_plan_path: Path) -> dict[str, dict]:
    plan = load_yaml_like(page_plan_path)
    if not isinstance(plan, dict):
        return {}
    allocations = plan.get("allocation") or []
    if not isinstance(allocations, list):
        return {}
    result = {}
    for item in allocations:
        if not isinstance(item, dict):
            continue
        section_id = str(item.get("section_id") or item.get("id") or "").strip()
        if section_id:
            result[section_id] = item
    return result


def scoring_rows_for_section(section: dict, scoring_matrix_path: Path) -> list[dict]:
    if not scoring_matrix_path.exists():
        return []
    rows = parse_markdown_table(read_text(scoring_matrix_path))
    declared = [normalize_title(item) for item in (section.get("scoring_items") or [])]
    matched = []
    for row in rows:
        item = first_present(
            row,
            [
                "Scoring item",
                "Score item",
                "评分项",
                "评审项目",
                "评分标准号",
                "建议标题",
                "Suggested title",
            ],
        )
        row_text = " ".join(str(value) for value in row.values())
        if any(item and item in normalize_title(row_text) for item in declared) or section_matches_text(section, row_text):
            matched.append(row)
    return matched


def mandatory_rows_for_section(section: dict, mandatory_matrix_path: Path) -> list[dict]:
    if not mandatory_matrix_path.exists():
        return []
    rows = parse_markdown_table(read_text(mandatory_matrix_path))
    req_ids = [normalize_title(req_id) for req_id in (section.get("req_ids") or [])]
    matched = []
    for row in rows:
        row_text = " ".join(str(value) for value in row.values())
        if any(req_id and req_id in normalize_title(row_text) for req_id in req_ids) or section_matches_text(section, row_text):
            matched.append(row)
    return matched


def score_from_row(row: dict) -> float:
    for value in row.values():
        match = re.search(r"\d+(?:\.\d+)?", str(value))
        if match:
            return float(match.group(0))
    return 0.0


def infer_contract_slots(section: dict, scoring_rows: list[dict], mandatory_rows: list[dict]) -> list[dict]:
    slots: list[dict] = [
        {
            "name": "requirement_response",
            "purpose": "Map this section to procurement requirements and state how the bid responds.",
            "source": "req_ids",
        },
        {
            "name": "scoring_response",
            "purpose": "Make scoring items easy for reviewers to find without creating new visible headings.",
            "source": "scoring_items",
        },
        {
            "name": "implementation_process",
            "purpose": "Describe concrete steps, responsibilities, controls, and sequence.",
            "source": "generic",
        },
        {
            "name": "deliverables",
            "purpose": "List outputs, documents, artifacts, and handover material tied to this section.",
            "source": "generic",
        },
        {
            "name": "acceptance_criteria",
            "purpose": "State how this content can be checked, accepted, or verified.",
            "source": "generic",
        },
    ]

    text = normalize_title(
        " ".join(
            [
                str(section.get("title") or ""),
                " ".join(str(item) for item in section.get("scoring_items") or []),
                " ".join(" ".join(str(value) for value in row.values()) for row in scoring_rows + mandatory_rows),
            ]
        )
    )
    conditional = [
        (("interface", "接口", "integration", "集成", "对接"), "interface_and_integration", "Describe boundaries, inputs/outputs, protocols, exceptions, and integration responsibilities."),
        (("data", "数据", "input", "output", "输入", "输出"), "input_output", "Describe input data, output results, data quality, and traceability."),
        (("deploy", "deployment", "部署", "环境"), "deployment", "Describe deployment topology, environment, release, and operational constraints."),
        (("nonfunctional", "performance", "security", "reliability", "性能", "安全", "可靠", "非功能"), "non_functional", "Describe performance, security, reliability, maintainability, and monitoring measures."),
        (("training", "培训", "support", "售后", "运维", "服务"), "service_and_training", "Describe service response, training, operation, maintenance, and support evidence."),
    ]
    existing = {slot["name"] for slot in slots}
    for needles, name, purpose in conditional:
        if name not in existing and any(normalize_title(needle) in text for needle in needles):
            slots.append({"name": name, "purpose": purpose, "source": "inferred_from_section_context"})
            existing.add(name)
    return slots


def section_content_weight(section: dict, scoring_rows: list[dict], mandatory_rows: list[dict], allocation: dict) -> dict:
    scoring = sum(score_from_row(row) for row in scoring_rows)
    if scoring <= 0 and section.get("scoring_items"):
        scoring = float(len(section.get("scoring_items") or []))
    mandatory = 3.0 * len(mandatory_rows) + 2.0 * len(section.get("req_ids") or [])
    pages_min = number_value(section.get("page_budget_min"))
    pages_max = number_value(section.get("page_budget_max"))
    if pages_min is None:
        pages_min = number_value(allocation.get("pages_min")) if allocation else None
    if pages_max is None:
        pages_max = number_value(allocation.get("pages_max")) if allocation else None
    complexity = max(0.0, ((pages_min or 0.0) + (pages_max or 0.0)) / 2.0)
    evidence = float(len(section.get("scoring_items") or []) + len(section.get("req_ids") or []))
    findability = 2.0 if section.get("scoring_items") else 0.0
    total = scoring + mandatory + complexity + evidence + findability
    return {
        "scoring": round(scoring, 2),
        "mandatory": round(mandatory, 2),
        "complexity": round(complexity, 2),
        "evidence": round(evidence, 2),
        "reviewer_findability": round(findability, 2),
        "total": round(total, 2),
        "reasons": [
            "scoring item score/count",
            "S0/S1 or req_ids coverage",
            "page budget complexity",
            "available section evidence",
            "reviewer findability for scoring items",
        ],
    }


def enrich_writing_plan_with_content_contracts(
    writing_plan_path: Path,
    page_plan_path: Path,
    scoring_matrix_path: Path,
    mandatory_matrix_path: Path,
    out_path: Path,
    chars_per_page: int,
) -> dict:
    plan = load_writing_plan(writing_plan_path)
    allocations = load_page_allocations(page_plan_path)
    enriched = deepcopy(plan)
    sections_enriched = 0

    for section in enriched.get("sections", []):
        if not isinstance(section, dict):
            continue
        section_id = section_id_of(section)
        allocation = allocations.get(section_id, {})
        scoring_rows = scoring_rows_for_section(section, scoring_matrix_path)
        mandatory_rows = mandatory_rows_for_section(section, mandatory_matrix_path)

        if number_value(section.get("page_budget_min")) is None and allocation:
            section["page_budget_min"] = allocation.get("pages_min", section.get("page_budget_min", ""))
        if number_value(section.get("page_budget_max")) is None and allocation:
            section["page_budget_max"] = allocation.get("pages_max", section.get("page_budget_max", ""))

        page_min = number_value(section.get("page_budget_min"))
        page_max = number_value(section.get("page_budget_max"))
        if page_min is not None:
            section["target_words_min"] = int(round(page_min * chars_per_page))
        if page_max is not None:
            section["target_words_max"] = int(round(page_max * chars_per_page))

        section["structure_locked"] = True
        section["content_weight"] = section_content_weight(section, scoring_rows, mandatory_rows, allocation)
        section["content_contract"] = {
            "hidden": True,
            "rules": [
                "do_not_create_visible_headings_from_slots",
                "do_not_change_section_title_order_or_level",
                "expand_only_with_confirmed_sources",
                "prefer_tables_processes_inputs_outputs_deliverables_acceptance",
                "forbid_padding_generic_slogans_and_repeated_rephrasing",
            ],
            "slots": infer_contract_slots(section, scoring_rows, mandatory_rows),
            "expansion_policy": {
                "preferred": [
                    "procurement requirement response",
                    "scoring item response",
                    "implementation process",
                    "inputs and outputs",
                    "interfaces and integration",
                    "deliverables and acceptance evidence",
                ],
                "forbidden": [
                    "generic guarantee sections",
                    "summary padding",
                    "repeated paraphrase",
                    "unsupported commitments",
                    "new visible headings from hidden slots",
                ],
            },
        }
        sections_enriched += 1

    write_yaml_like(out_path, enriched)
    return {"out_path": str(out_path), "sections_enriched": sections_enriched}


def load_yaml_like(path: Path) -> object:
    if not path.exists():
        return {}
    text = read_text(path)
    if yaml is not None:
        return yaml.safe_load(text) or {}
    try:
        return json.loads(text)
    except Exception:
        return {}


def write_yaml_like(path: Path, obj: object) -> None:
    ensure_parent(path)
    if yaml is not None:
        path.write_text(yaml.safe_dump(obj, allow_unicode=True, sort_keys=False), encoding="utf-8")
    else:
        path.write_text(json.dumps(obj, ensure_ascii=False, indent=2), encoding="utf-8")


def now_iso() -> str:
    return datetime.now().isoformat(timespec="seconds")


def safe_section_id(section_id: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", section_id.strip())
    cleaned = cleaned.strip(". ")
    return cleaned or "section"


def load_writing_plan(path: Path) -> dict:
    plan = load_yaml_like(path)
    if not isinstance(plan, dict):
        raise RuntimeError(f"writing plan is invalid or missing: {path}")
    sections = plan.get("sections")
    if isinstance(sections, dict):
        normalized_sections = []
        for key, value in sections.items():
            if not isinstance(value, dict):
                raise RuntimeError(f"writing plan section {key} must be an object: {path}")
            section = dict(value)
            section.setdefault("id", str(key))
            normalized_sections.append(section)
        plan["sections"] = normalized_sections
    elif not isinstance(sections, list):
        raise RuntimeError(f"writing plan must contain a sections list: {path}")
    return plan


def section_id_of(section: dict) -> str:
    return str(section.get("id") or section.get("section_id") or "").strip()


def find_section(plan: dict, section_id: str) -> Optional[dict]:
    for section in plan.get("sections", []):
        if isinstance(section, dict) and section_id_of(section) == section_id:
            return section
    return None


def section_file_path(section: dict, sections_dir: Path) -> Path:
    raw = str(section.get("file_path") or "").strip()
    if raw:
        return Path(raw)
    return sections_dir / f"{safe_section_id(section_id_of(section))}.md"


def lock_file_path(section_id: str, locks_dir: Path) -> Path:
    return locks_dir / f"{safe_section_id(section_id)}.lock"


def section_manual_flags(section: dict) -> list:
    flags = section.get("manual_flags") or []
    return flags if isinstance(flags, list) else [flags]


def section_dependencies_done(plan: dict, section: dict) -> bool:
    deps = section.get("depends_on") or []
    if not isinstance(deps, list):
        deps = [deps]
    for dep in deps:
        dep_id = str(dep).strip()
        if not dep_id:
            continue
        dep_section = find_section(plan, dep_id)
        if dep_section is None or str(dep_section.get("status") or "").strip() != "drafted":
            return False
    return True


def section_sort_key(index_and_section: tuple[int, dict]) -> tuple[float, int]:
    index, section = index_and_section
    order = number_value(section.get("merge_order"))
    if order is None:
        order = number_value(section_id_of(section).replace(".", ""))
    return (order if order is not None else float(index), index)


def section_budget_words(section: dict, chars_per_page: int) -> tuple[Optional[int], Optional[int], list[str]]:
    sources = []
    min_words = number_value(section.get("target_words_min"))
    max_words = number_value(section.get("target_words_max"))
    if min_words is not None:
        sources.append("target_words_min")
    if max_words is not None:
        sources.append("target_words_max")

    if min_words is None:
        page_min = number_value(section.get("page_budget_min"))
        if page_min is not None:
            min_words = page_min * chars_per_page
            sources.append("page_budget_min")
    if max_words is None:
        page_max = number_value(section.get("page_budget_max"))
        if page_max is not None:
            max_words = page_max * chars_per_page
            sources.append("page_budget_max")

    min_budget = int(round(min_words)) if min_words is not None else None
    max_budget = int(round(max_words)) if max_words is not None else None
    return min_budget, max_budget, sources


def section_budget_report(
    section: dict,
    drafted_words: int,
    chars_per_page: int,
    min_ratio: float,
    max_ratio: float,
) -> dict:
    min_budget, max_budget, sources = section_budget_words(section, chars_per_page)
    section_id = section_id_of(section)
    title = str(section.get("title") or "")
    result = "NOT_SET"
    issues = []
    warnings = []
    lower_required = None
    upper_allowed = None

    if min_budget is not None or max_budget is not None:
        result = "PASS"
    if min_budget is not None:
        lower_required = int(math.ceil(min_budget * min_ratio))
        if drafted_words < lower_required:
            result = "FAIL"
            issues.append(
                f"{section_id} {title}: drafted_words={drafted_words} is below "
                f"{lower_required} ({min_ratio:.0%} of min budget {min_budget})."
            )
    if max_budget is not None:
        upper_allowed = int(math.floor(max_budget * max_ratio))
        if drafted_words > upper_allowed:
            if result != "FAIL":
                result = "WARN"
            warnings.append(
                f"{section_id} {title}: drafted_words={drafted_words} is above "
                f"{upper_allowed} ({max_ratio:.0%} of max budget {max_budget}); "
                "review for repetition, filler, or unsupported content, but do not block unless the user set a hard page limit."
            )

    return {
        "section_id": section_id,
        "title": title,
        "result": result,
        "drafted_words": drafted_words,
        "target_words_min": min_budget,
        "target_words_max": max_budget,
        "lower_required": lower_required,
        "upper_allowed": upper_allowed,
        "sources": sources,
        "issues": issues,
        "warnings": warnings,
    }


def build_expansion_tasks(section: dict, drafted_words: int, chars_per_page: int) -> list[str]:
    """Build Expansion Tasks list for a section that failed budget check."""
    blueprint = section_writing_blueprint(section, chars_per_page)
    if not blueprint:
        return []
    min_budget, _, _ = section_budget_words(section, chars_per_page)
    words_to_lower = max((min_budget or 0) - drafted_words, 0) if min_budget else 0
    if not words_to_lower:
        return [block["label"] for block in blueprint[:6]]
    candidate_blocks = blueprint[: max(1, min(len(blueprint), 6))]
    per_block = int(math.ceil(words_to_lower / max(len(candidate_blocks), 1)))
    tasks = []
    for block in candidate_blocks:
        tasks.append(
            f"继续扩写 {block['label']}（{block['name']}）约 {per_block} 字：{block['guidance']}"
        )
    return tasks


def section_budget_failure_message(
    report: dict,
    expansion_tasks: Optional[list[str]] = None,
    section: Optional[dict] = None,
    drafted_words: int = 0,
    chars_per_page: int = DEFAULT_CHARS_PER_PAGE_ESTIMATE,
) -> str:
    guidance = (
        "Expand or trim this section under its own scoring items, procurement requirements, "
        "processes, inputs/outputs, deliverables, acceptance criteria, interfaces, deployment, "
        "or non-functional constraints. If the confirmed source material cannot support the "
        "budget, complete the section as blocked or MANUAL with notes instead of marking it drafted."
    )
    msg = "section page/word budget check failed:\n- " + "\n- ".join(report["issues"]) + "\n" + guidance
    if expansion_tasks:
        task_lines = "\n".join(f"  {i+1}. [{task.split('（')[1] if '（' in task else task}] {task.split('）：')[1] if '）：' in task else task}"
                               for i, task in enumerate(expansion_tasks))
        msg = (
            f"BUDGET CHECK FAIL\n"
            f"  当前字数：{drafted_words}字\n"
            f"  下限要求：{report.get('lower_required', '?')}字\n"
            f"  差距：{max((report.get('lower_required') or 0) - drafted_words, 0)}字\n\n"
            f"Expansion Tasks:\n"
            + task_lines
            + f"\n\n操作：在原文件末尾补充上述内容，然后重新提交：\n"
            f"  python tenderctl.py complete-section --section-id {report.get('section_id', '')} --status drafted"
        )
    return msg


def append_manual_decision(path: Path, section_id: str, title: str, reason: str, suggestions: list[str]) -> None:
    existing = read_text(path)
    lines = []
    if existing.strip():
        lines.append(existing.rstrip())
        lines.append("")
    lines.extend(
        [
            f"## {now_iso()} {section_id} {title}".rstrip(),
            "",
            f"- 状态：needs_expansion",
            f"- 原因：{reason}",
            "- 建议扩写点：",
        ]
    )
    if suggestions:
        lines.extend([f"  - {item}" for item in suggestions])
    else:
        lines.append("  - 围绕评分项响应、采购需求响应、业务流程、输入输出、接口边界、部署位置、异常处理、交付物和验收依据补足。")
    write_text(path, "\n".join(lines) + "\n")


def status_sections_from_plan(plan: dict, locks_dir: Path, sections_dir: Path) -> list[dict]:
    items = []
    for section in plan.get("sections", []):
        if not isinstance(section, dict):
            continue
        section_id = section_id_of(section)
        if not section_id:
            continue
        lock_path = lock_file_path(section_id, locks_dir)
        file_path = section_file_path(section, sections_dir)
        items.append(
            {
                "id": section_id,
                "title": section.get("title", ""),
                "status": section.get("status", "pending"),
                "owner": section.get("owner", ""),
                "started_at": section.get("started_at", ""),
                "completed_at": section.get("completed_at", ""),
                "drafted_words": section.get("drafted_words", 0),
                "file_path": str(file_path),
                "lock_path": str(lock_path) if lock_path.exists() else "",
                "notes": section.get("notes", ""),
                "manual_flags": section_manual_flags(section),
                "target_words_min": section.get("target_words_min", ""),
                "target_words_max": section.get("target_words_max", ""),
                "page_budget_min": section.get("page_budget_min", ""),
                "page_budget_max": section.get("page_budget_max", ""),
            }
        )
    return items


def write_section_status(plan: dict, status_yml: Path, status_md: Path, locks_dir: Path, sections_dir: Path) -> None:
    sections = status_sections_from_plan(plan, locks_dir, sections_dir)
    status_obj = {"updated_at": now_iso(), "sections": sections}
    write_yaml_like(status_yml, status_obj)

    lines = [
        "# 章节写作状态",
        "",
        "| ID | 标题 | 状态 | 负责人 | 字数 | 文件 | 备注 |",
        "|---|---|---|---|---:|---|---|",
    ]
    for item in sections:
        lines.append(
            "| {id} | {title} | {status} | {owner} | {drafted_words} | {file_path} | {notes} |".format(
                id=item["id"],
                title=str(item["title"]).replace("|", "\\|"),
                status=item["status"],
                owner=str(item["owner"]).replace("|", "\\|"),
                drafted_words=item["drafted_words"] or 0,
                file_path=item["file_path"],
                notes=str(item["notes"]).replace("|", "\\|"),
            )
        )
    write_text(status_md, "\n".join(lines) + "\n")


def section_brief_path(section_id: str, briefs_dir: Path) -> Path:
    return briefs_dir / f"{section_id}.md"


def list_value_text(value: object) -> str:
    if isinstance(value, list):
        items = [str(item).strip() for item in value if str(item).strip()]
        return ", ".join(items) if items else "not set"
    text = str(value or "").strip()
    return text if text else "not set"


def contract_slot_name(slot: object) -> str:
    if isinstance(slot, dict):
        return str(slot.get("name") or slot.get("slot") or "").strip()
    return str(slot or "").strip()


SECTION_BLOCK_RULES = {
    "requirement_response": {
        "label": "采购需求响应",
        "keywords": ["需求", "采购", "响应", "要求", "满足", "对应"],
        "guidance": "把招标要求转成本节的响应动作、边界、交付和不承诺事项。",
    },
    "scoring_response": {
        "label": "评分项响应",
        "keywords": ["评分", "得分", "评审", "分值", "指标", "题目"],
        "guidance": "逐条承接本节评分项，让专家能看到本节为什么能得分。",
    },
    "implementation_process": {
        "label": "实施流程",
        "keywords": ["流程", "步骤", "阶段", "执行", "处理", "调度", "发布"],
        "guidance": "写清步骤、责任、输入、输出、异常处理和控制点。",
    },
    "input_output": {
        "label": "输入输出",
        "keywords": ["输入", "输出", "字段", "格式", "数据项", "结果", "参数"],
        "guidance": "列明数据、接口、文档、配置、模型结果或验收材料的流转。",
    },
    "interface_and_integration": {
        "label": "接口与集成",
        "keywords": ["接口", "API", "协议", "集成", "对接", "调用", "异常"],
        "guidance": "说明接口边界、调用方式、输入输出、异常码和对接责任。",
    },
    "deployment": {
        "label": "部署与环境",
        "keywords": ["部署", "环境", "容器", "服务器", "网络", "发布", "配置"],
        "guidance": "说明部署拓扑、运行环境、发布流程和运维边界。",
    },
    "non_functional": {
        "label": "非功能约束",
        "keywords": ["性能", "安全", "可靠", "可用", "并发", "监控", "日志"],
        "guidance": "写清性能、安全、可靠性、可观测性和维护措施。",
    },
    "deliverables": {
        "label": "交付物",
        "keywords": ["交付", "文档", "清单", "报告", "源码", "部署包", "手册"],
        "guidance": "写清交付物名称、内容、提交时点、确认方式和关联验收。",
    },
    "acceptance_criteria": {
        "label": "验收依据",
        "keywords": ["验收", "测试", "确认", "标准", "材料", "证明", "通过"],
        "guidance": "写清验收材料、验收动作、判定标准和风险处理。",
    },
    "service_and_training": {
        "label": "服务与培训",
        "keywords": ["培训", "服务", "运维", "售后", "响应", "支持", "教材"],
        "guidance": "说明培训对象、内容、形式、服务响应和运维支撑证据。",
    },
}


def section_blueprint_slots(section: dict) -> list[str]:
    contract = section.get("content_contract") if isinstance(section.get("content_contract"), dict) else {}
    slots = [contract_slot_name(slot) for slot in contract.get("slots", [])] if contract else []
    ordered = []
    for name in [
        "scoring_response",
        "requirement_response",
        "implementation_process",
        "input_output",
        "interface_and_integration",
        "deployment",
        "non_functional",
        "deliverables",
        "acceptance_criteria",
        "service_and_training",
    ]:
        if name in slots or name in {"scoring_response", "implementation_process", "deliverables", "acceptance_criteria"}:
            ordered.append(name)
    for name in slots:
        if name and name not in ordered:
            ordered.append(name)
    return ordered


def section_writing_blueprint(section: dict, chars_per_page: int = DEFAULT_CHARS_PER_PAGE_ESTIMATE) -> list[dict]:
    min_words, max_words, _ = section_budget_words(section, chars_per_page)
    target = min_words or max_words or 0
    slots = section_blueprint_slots(section)
    if not slots:
        return []
    weights = {}
    for name in slots:
        if name == "scoring_response":
            weights[name] = 1.8 if section.get("scoring_items") else 1.0
        elif name == "requirement_response":
            weights[name] = 1.5 if section.get("req_ids") else 1.0
        elif name in {"implementation_process", "input_output", "interface_and_integration"}:
            weights[name] = 1.4
        elif name in {"deliverables", "acceptance_criteria"}:
            weights[name] = 1.1
        else:
            weights[name] = 1.0
    total_weight = sum(weights.values()) or 1.0
    blocks = []
    allocated = 0
    for index, name in enumerate(slots, start=1):
        rule = SECTION_BLOCK_RULES.get(name, {"label": name, "guidance": "按本节材料自然展开。"})
        if target:
            target_words = int(round(target * weights[name] / total_weight))
            allocated += target_words
        else:
            target_words = 0
        blocks.append(
            {
                "order": index,
                "name": name,
                "label": rule["label"],
                "target_words": target_words,
                "guidance": rule["guidance"],
            }
        )
    if target and blocks:
        blocks[-1]["target_words"] += target - allocated
    return blocks


def text_has_block(text: str, block_name: str, section: Optional[dict] = None) -> bool:
    normalized = normalize_title(text)
    if block_name == "scoring_response" and section:
        items = [normalize_title(item) for item in section.get("scoring_items") or []]
        if items and all(item and item in normalized for item in items):
            return True
    rule = SECTION_BLOCK_RULES.get(block_name)
    if not rule:
        return True
    return any(normalize_title(keyword) in normalized for keyword in rule["keywords"])


def section_gap_report(section: dict, markdown: str, chars_per_page: int, min_ratio: float = 0.90) -> dict:
    drafted_words = effective_char_count(strip_markdown_for_count(markdown))
    budget = section_budget_report(section, drafted_words, chars_per_page, min_ratio, 1.10)
    blueprint = section_writing_blueprint(section, chars_per_page)
    missing_blocks = [block for block in blueprint if not text_has_block(markdown, block["name"], section)]
    words_to_lower = max((budget.get("lower_required") or 0) - drafted_words, 0)
    result = "PASS"
    if budget["result"] == "FAIL" or missing_blocks:
        result = "FAIL"
    expansion_tasks = []
    if words_to_lower:
        candidate_blocks = missing_blocks or blueprint
        candidate_blocks = candidate_blocks[: max(1, min(len(candidate_blocks), 6))]
        per_block = int(math.ceil(words_to_lower / max(len(candidate_blocks), 1)))
        for block in candidate_blocks:
            expansion_tasks.append(
                f"继续扩写 {block['label']}（{block['name']}）约 {per_block} 字：{block['guidance']}"
            )
    for block in missing_blocks:
        task = f"补齐 {block['label']}（{block['name']}）：{block['guidance']}"
        if task not in expansion_tasks:
            expansion_tasks.append(task)
    return {
        "section_id": section_id_of(section),
        "title": section.get("title", ""),
        "result": result,
        "drafted_words": drafted_words,
        "budget": budget,
        "words_to_lower_required": words_to_lower,
        "blueprint": blueprint,
        "missing_blocks": missing_blocks,
        "expansion_tasks": expansion_tasks,
    }


def write_section_brief(section: dict, brief_path: Path, file_path: Path) -> None:
    section_id = section_id_of(section)
    title = str(section.get("title") or "").strip()
    contract = section.get("content_contract") if isinstance(section.get("content_contract"), dict) else {}
    slots = [contract_slot_name(slot) for slot in contract.get("slots", [])] if contract else []
    slots = [slot for slot in slots if slot]
    weights = section.get("content_weight") if isinstance(section.get("content_weight"), dict) else {}

    page_min = section.get("page_budget_min", "")
    page_max = section.get("page_budget_max", "")
    words_min = section.get("target_words_min", "")
    words_max = section.get("target_words_max", "")
    page_text = f"{page_min}-{page_max} pages" if page_min or page_max else "not set"
    words_text = f"{words_min}-{words_max}" if words_min or words_max else "not set"

    lines = [
        f"# Section Brief: {section_id} {title}".rstrip(),
        "",
        "This is an internal writing brief for the agent. Do not copy it into the tender body.",
        "",
        "## Budget",
        "",
        f"- Page budget: {page_text}",
        f"- Target words: {words_text}",
        f"- Draft file: {file_path}",
        "",
        "## Coverage",
        "",
        f"- Requirement IDs: {list_value_text(section.get('req_ids'))}",
        f"- Scoring items: {list_value_text(section.get('scoring_items'))}",
        "",
        "## Content Weights",
        "",
    ]
    if weights:
        for key, value in weights.items():
            lines.append(f"- {key}: {value}")
    else:
        lines.append("- not set")

    lines.extend(["", "## Hidden Slots", ""])
    if slots:
        lines.extend([f"- {slot}" for slot in slots])
    else:
        lines.append("- not set")

    blueprint = section_writing_blueprint(section)
    lines.extend(["", "## Writing Blueprint", ""])
    if blueprint:
        lines.append("| Order | Block | Target words | Guidance |")
        lines.append("|---:|---|---:|---|")
        for block in blueprint:
            target_words = block["target_words"] or ""
            lines.append(
                f"| {block['order']} | {block['label']} (`{block['name']}`) | {target_words} | {block['guidance']} |"
            )
    else:
        lines.append("- not set")

    lines.extend(
        [
            "",
            "## Drafting Instructions",
            "",
            "- Write the internal micro-outline from the slots before drafting.",
            "- Expand first on scoring items, mandatory requirements, process, inputs, outputs, interfaces, deliverables, and acceptance evidence.",
            "- Keep the visible section title and heading level unchanged.",
            "- Do not create visible headings from hidden slots.",
            "- If confirmed sources cannot support the budget, mark the section blocked or MANUAL instead of padding.",
            "- After the first draft, run section-gap or complete-section; if it fails, continue the listed expansion tasks before merging.",
        ]
    )
    write_text(brief_path, "\n".join(lines) + "\n")


def cmd_claim_section(args) -> int:
    plan_path = Path(args.writing_plan)
    locks_dir = Path(args.locks_dir)
    sections_dir = Path(args.sections_dir)
    briefs_dir = Path(getattr(args, "section_briefs_dir", "work/40_drafts/section_briefs"))
    status_yml = Path(args.section_status)
    status_md = Path(args.out_md)
    owner = args.owner or "agent"

    plan = load_writing_plan(plan_path)
    candidate = None
    if args.section_id:
        candidate = find_section(plan, args.section_id)
        if candidate is None:
            raise RuntimeError(f"section not found: {args.section_id}")
    else:
        ordered = sorted(enumerate(plan.get("sections", [])), key=section_sort_key)
        for _, section in ordered:
            if not isinstance(section, dict):
                continue
            status = str(section.get("status") or "pending").strip()
            if status not in {"pending", "needs_revision", "needs_expansion"}:
                continue
            if section_manual_flags(section):
                continue
            if not section_dependencies_done(plan, section):
                continue
            if section_file_path(section, sections_dir).exists():
                continue
            candidate = section
            break
        if candidate is None:
            raise RuntimeError("no claimable section found")

    section_id = section_id_of(candidate)
    if not section_id:
        raise RuntimeError("selected section has no id")
    status = str(candidate.get("status") or "pending").strip()
    if status not in {"pending", "needs_revision", "needs_expansion"}:
        raise RuntimeError(f"section {section_id} is not claimable; status={status}")
    if section_manual_flags(candidate):
        raise RuntimeError(f"section {section_id} has manual flags and cannot be claimed")
    if not section_dependencies_done(plan, candidate):
        raise RuntimeError(f"section {section_id} dependencies are not drafted")

    file_path = section_file_path(candidate, sections_dir)
    if file_path.exists() and not args.allow_existing_file:
        raise RuntimeError(f"section file already exists: {file_path}")
    brief_path = section_brief_path(section_id, briefs_dir)

    locks_dir.mkdir(parents=True, exist_ok=True)
    lock_path = lock_file_path(section_id, locks_dir)
    started_at = now_iso()
    lock_payload = {
        "section_id": section_id,
        "title": candidate.get("title", ""),
        "owner": owner,
        "started_at": started_at,
        "file_path": str(file_path),
        "section_brief_path": str(brief_path),
    }
    try:
        with lock_path.open("x", encoding="utf-8") as f:
            json.dump(lock_payload, f, ensure_ascii=False, indent=2)
    except FileExistsError as exc:
        raise RuntimeError(f"section already claimed: {section_id} ({lock_path})") from exc

    candidate["status"] = "in_progress"
    candidate["owner"] = owner
    candidate["started_at"] = started_at
    candidate["lock_path"] = str(lock_path)
    candidate["file_path"] = str(file_path)
    candidate["section_brief_path"] = str(brief_path)
    candidate["updated_at"] = started_at
    if not getattr(args, "skip_section_brief", False):
        write_section_brief(candidate, brief_path, file_path)
    write_yaml_like(plan_path, plan)
    write_section_status(plan, status_yml, status_md, locks_dir, sections_dir)

    print(json.dumps(lock_payload, ensure_ascii=False, indent=2))
    return 0


def cmd_complete_section(args) -> int:
    plan_path = Path(args.writing_plan)
    locks_dir = Path(args.locks_dir)
    sections_dir = Path(args.sections_dir)
    status_yml = Path(args.section_status)
    status_md = Path(args.out_md)
    completed_status = args.status

    plan = load_writing_plan(plan_path)
    section = find_section(plan, args.section_id)
    if section is None:
        raise RuntimeError(f"section not found: {args.section_id}")
    section_id = section_id_of(section)
    lock_path = lock_file_path(section_id, locks_dir)
    if lock_path.exists():
        try:
            lock_payload = json.loads(read_text(lock_path))
        except Exception:
            lock_payload = {}
        lock_owner = str(lock_payload.get("owner") or "")
        if args.owner and lock_owner and lock_owner != args.owner:
            raise RuntimeError(f"section {section_id} is locked by {lock_owner}, not {args.owner}")
    elif not args.allow_no_lock:
        raise RuntimeError(f"section {section_id} has no lock; use claim-section first")

    file_path = Path(args.section_file) if args.section_file else section_file_path(section, sections_dir)
    if completed_status == "drafted":
        if not file_path.exists():
            raise RuntimeError(f"section file not found: {file_path}")
        if not read_text(file_path).strip():
            raise RuntimeError(f"section file is empty: {file_path}")

    completed_at = now_iso()
    drafted_words = int(args.words) if args.words else 0
    if completed_status == "drafted" and not drafted_words:
        drafted_words = effective_char_count(strip_markdown_for_count(read_text(file_path)))
    if completed_status == "drafted":
        chars = resolve_chars_per_page_from_args(args)
        budget_report = section_budget_report(
            section,
            drafted_words,
            chars,
            float(args.min_budget_ratio),
            float(args.max_budget_ratio),
        )
        if budget_report["result"] == "FAIL":
            expansion = build_expansion_tasks(section, drafted_words, chars)
            completed_at = now_iso()
            section["status"] = "needs_expansion"
            section["completed_at"] = completed_at
            section["updated_at"] = completed_at
            section["file_path"] = str(file_path)
            section["drafted_words"] = drafted_words
            if args.owner:
                section["owner"] = args.owner
            section["notes"] = (
                f"目标字数 {budget_report.get('target_words_min')}，当前字数 {drafted_words}，"
                f"缺口 {max((budget_report.get('lower_required') or 0) - drafted_words, 0)}。"
            )
            flags = section_manual_flags(section)
            if "needs_expansion: below 90% of target_words_min" not in flags:
                flags.append("needs_expansion: below 90% of target_words_min")
            section["manual_flags"] = flags
            section.pop("lock_path", None)
            if lock_path.exists() and not args.keep_lock:
                lock_path.unlink()
            write_yaml_like(plan_path, plan)
            write_section_status(plan, status_yml, status_md, locks_dir, sections_dir)
            append_manual_decision(
                Path(args.manual_decisions),
                section_id,
                str(section.get("title") or ""),
                section["notes"],
                expansion,
            )
            raise RuntimeError(
                section_budget_failure_message(budget_report, expansion, section, drafted_words, chars)
            )

    section["status"] = completed_status
    section["completed_at"] = completed_at
    section["updated_at"] = completed_at
    section["file_path"] = str(file_path)
    section["drafted_words"] = drafted_words
    if args.owner:
        section["owner"] = args.owner
    if args.notes:
        section["notes"] = args.notes
    if completed_status in {"blocked", "MANUAL", "needs_revision", "needs_expansion"} and args.notes:
        flags = section_manual_flags(section)
        if args.notes not in flags:
            flags.append(args.notes)
        section["manual_flags"] = flags
    section.pop("lock_path", None)

    if lock_path.exists() and not args.keep_lock:
        lock_path.unlink()

    write_yaml_like(plan_path, plan)
    write_section_status(plan, status_yml, status_md, locks_dir, sections_dir)
    print(f"[OK] section {section_id} -> {completed_status}")
    return 0


def cmd_merge_sections(args) -> int:
    plan_path = Path(args.writing_plan)
    sections_dir = Path(args.sections_dir)
    locks_dir = Path(args.locks_dir)
    status_yml = Path(args.section_status)
    status_md = Path(args.out_md)
    out_draft = Path(args.out_draft)

    plan = load_writing_plan(plan_path)
    manifest = read_manifest(Path(getattr(args, "manifest", "work/00_manifest.yml")))
    chars_per_page = resolve_chars_per_page(args.chars_per_page, manifest)
    blockers = []
    parts = []
    actual_words_by_section = {}
    target_words_min_by_section = {}
    ordered = sorted(enumerate(plan.get("sections", [])), key=section_sort_key)
    for _, section in ordered:
        if not isinstance(section, dict):
            continue
        section_id = section_id_of(section)
        if not section_id:
            continue
        status = str(section.get("status") or "pending").strip()
        file_path = section_file_path(section, sections_dir)
        if status == "drafted":
            if not file_path.exists():
                blockers.append(f"{section_id}: status=drafted but file missing: {file_path}")
                continue
            text = read_text(file_path).strip()
            if not text:
                blockers.append(f"{section_id}: section file is empty: {file_path}")
                continue
            drafted_words = effective_char_count(strip_markdown_for_count(text))
            min_words, _, _ = section_budget_words(section, chars_per_page)
            if min_words is not None:
                target_words_min_by_section[section_id] = min_words
            actual_words_by_section[section_id] = drafted_words
            budget_report = section_budget_report(
                section,
                drafted_words,
                chars_per_page,
                float(args.min_budget_ratio),
                float(args.max_budget_ratio),
            )
            if budget_report["result"] == "FAIL":
                blockers.append(section_budget_failure_message(budget_report, None, section, drafted_words, chars_per_page))
                continue
            section["drafted_words"] = drafted_words
            parts.append(text)
        elif status in {"pending", "in_progress", "blocked", "MANUAL", "needs_revision"}:
            message = f"{section_id}: status={status}"
            if args.allow_partial:
                continue
            blockers.append(message)

    target_min_total = sum(target_words_min_by_section.values())
    actual_total = sum(actual_words_by_section.values())
    lower_required_total = int(math.ceil(target_min_total * float(args.min_budget_ratio))) if target_min_total else 0
    target_pages_min = int(manifest.get("target_pages_min") or 0)
    draft_lower_required = int(math.ceil(target_pages_min * chars_per_page * float(args.min_budget_ratio))) if target_pages_min else 0
    if not args.allow_partial:
        total_issues = []
        if lower_required_total and actual_total < lower_required_total:
            total_issues.append(
                f"sum(actual_words_by_section)={actual_total} < {lower_required_total} "
                f"(90% of sum(target_words_min_by_section)={target_min_total})"
            )
        if draft_lower_required and actual_total < draft_lower_required:
            total_issues.append(
                f"v1.md actual_words={actual_total} < {draft_lower_required} "
                f"(90% of target_pages_min={target_pages_min} * chars_per_page={chars_per_page})"
            )
        if total_issues:
            short_sections = []
            for section_id, target in target_words_min_by_section.items():
                actual = actual_words_by_section.get(section_id, 0)
                lower = int(math.ceil(target * float(args.min_budget_ratio)))
                if actual < lower:
                    short_sections.append(f"{section_id}: actual={actual}, lower_required={lower}, target_words_min={target}")
            blockers.insert(
                0,
                "total draft word budget check failed:\n  - "
                + "\n  - ".join(total_issues)
                + ("\n  - short sections: " + "; ".join(short_sections) if short_sections else ""),
            )

    if blockers:
        raise RuntimeError("cannot merge sections:\n- " + "\n- ".join(blockers))

    write_text(out_draft, "\n\n".join(parts).rstrip() + "\n")
    merged_at = now_iso()
    for section in plan.get("sections", []):
        if isinstance(section, dict) and str(section.get("status") or "") == "drafted":
            section["merged_at"] = merged_at
    write_yaml_like(plan_path, plan)
    write_section_status(plan, status_yml, status_md, locks_dir, sections_dir)
    print(f"[OK] merged {len(parts)} sections into: {out_draft}")
    return 0


def cmd_normalize_writing_plan(args) -> int:
    plan_path = Path(args.writing_plan)
    plan = load_writing_plan(plan_path)
    write_yaml_like(plan_path, plan)
    section_count = len([section for section in plan.get("sections", []) if isinstance(section, dict)])
    print(f"[OK] normalized writing plan sections to list: {plan_path} ({section_count} sections)")
    return 0


def cmd_content_contract(args) -> int:
    manifest = read_manifest(Path(args.manifest))
    page_plan_report = audit_page_plan(Path(args.page_plan), manifest)
    if page_plan_report["result"] == "FAIL":
        issues = page_plan_report.get("issues") or ["page_plan.yml failed validation."]
        raise RuntimeError("页数预算未通过统筹检查；请回到 outline 阶段调整 page_plan.yml 和 writing_plan.yml。\n- " + "\n- ".join(issues))

    result = enrich_writing_plan_with_content_contracts(
        writing_plan_path=Path(args.writing_plan),
        page_plan_path=Path(args.page_plan),
        scoring_matrix_path=Path(args.scoring_matrix),
        mandatory_matrix_path=Path(args.mandatory_matrix),
        out_path=Path(args.out_writing_plan or args.writing_plan),
        chars_per_page=resolve_chars_per_page_from_args(args),
    )
    print(f"[OK] enriched {result['sections_enriched']} sections with hidden content contracts")
    print(f"[OK] wrote: {result['out_path']}")
    return 0


def cmd_draft_section(args) -> int:
    """
    End-to-end section drafting with automatic budget enforcement loop.

    Loop: generate content -> check budget -> if FAIL, generate expansion tasks
          -> re-prompt LLM with specific tasks -> repeat until PASS or max_cycles
    """
    manifest = read_manifest(Path(args.manifest))
    chars_per_page = resolve_chars_per_page(args.chars_per_page, manifest)
    plan_path = Path(args.writing_plan)
    plan = load_writing_plan(plan_path)
    sections_dir = Path(args.sections_dir)
    briefs_dir = Path(args.section_briefs_dir)
    status_yml = Path(args.section_status)
    status_md = Path(args.out_md)
    locks_dir = Path(args.locks_dir)

    # Select or validate section
    if args.section_id:
        section = find_section(plan, args.section_id)
        if section is None:
            raise RuntimeError(f"section not found: {args.section_id}")
    else:
        ordered = sorted(enumerate(plan.get("sections", [])), key=section_sort_key)
        for _, sec in ordered:
            if not isinstance(sec, dict):
                continue
            status = str(sec.get("status") or "pending").strip()
            if status not in {"pending", "needs_revision", "needs_expansion"}:
                continue
            if section_manual_flags(sec):
                continue
            if not section_dependencies_done(plan, sec):
                continue
            if section_file_path(sec, sections_dir).exists() and not args.allow_existing_file:
                continue
            section = sec
            break
        if section is None:
            raise RuntimeError("no claimable section found")

    section_id = section_id_of(section)
    if not section_id:
        raise RuntimeError("section has no id")

    file_path = section_file_path(section, sections_dir)
    brief_path = section_brief_path(section_id, briefs_dir)

    # If file exists and not allow_existing, fail
    if file_path.exists() and not args.allow_existing_file and not args.continue_existing:
        raise RuntimeError(f"section file already exists: {file_path}; use --continue-existing to continue")

    owner = args.owner or "agent"
    started_at = now_iso()

    # Create lock
    locks_dir.mkdir(parents=True, exist_ok=True)
    lock_path = lock_file_path(section_id, locks_dir)
    lock_payload = {
        "section_id": section_id,
        "title": section.get("title", ""),
        "owner": owner,
        "started_at": started_at,
        "file_path": str(file_path),
        "section_brief_path": str(brief_path),
    }
    try:
        with lock_path.open("x", encoding="utf-8") as f:
            json.dump(lock_payload, f, ensure_ascii=False, indent=2)
    except FileExistsError:
        if not args.allow_existing_file:
            raise RuntimeError(f"section already claimed: {section_id}")

    section["status"] = "in_progress"
    section["owner"] = owner
    section["started_at"] = started_at
    section["file_path"] = str(file_path)
    section["section_brief_path"] = str(brief_path)
    section["updated_at"] = started_at

    # Generate or load existing brief
    if not brief_path.exists():
        write_section_brief(section, brief_path, file_path)

    brief_content = read_text(brief_path)
    blueprint = section_writing_blueprint(section, chars_per_page)
    min_budget, max_budget, _ = section_budget_words(section, chars_per_page)

    # Determine target words
    target_min = min_budget or (max_budget if max_budget else 0)

    # Load existing content if continuing
    existing_content = ""
    if args.continue_existing and file_path.exists():
        existing_content = read_text(file_path)

    # Build system prompt for LLM drafting
    system_prompt = _build_draft_section_system_prompt()

    # Expansion history to avoid repeating the same instructions
    expansion_history = []
    max_cycles = int(args.max_cycles) if args.max_cycles else 3
    cycle = 0

    while cycle < max_cycles:
        cycle += 1
        print(f"[draft-section] cycle {cycle}/{max_cycles} for section {section_id}")

        # Build user prompt
        user_prompt = _build_draft_section_user_prompt(
            section_id=section_id,
            section_title=section.get("title", ""),
            brief_content=brief_content,
            blueprint=blueprint,
            target_min=target_min,
            chars_per_page=chars_per_page,
            existing_content=existing_content,
            expansion_history=expansion_history,
        )

        # Call LLM
        content = _call_llm_generate(system_prompt, user_prompt, args.model)

        if not content or not content.strip():
            raise RuntimeError(f"LLM returned empty content on cycle {cycle}")

        # Update existing content for next cycle
        if existing_content:
            # Append new content to existing
            existing_content = existing_content.rstrip() + "\n\n" + content
        else:
            existing_content = content

        # Write current result to section file (in case of failure, partial work is saved)
        write_text(file_path, existing_content)

        # Check budget
        drafted_words = effective_char_count(strip_markdown_for_count(existing_content))
        budget_report = section_budget_report(
            section, drafted_words, chars_per_page,
            float(args.min_budget_ratio), float(args.max_budget_ratio)
        )

        print(f"[draft-section] drafted_words={drafted_words}, target_min={target_min}, budget_result={budget_report['result']}")

        if budget_report["result"] != "FAIL":
            # PASS or WARN - we're good
            if budget_report["result"] == "WARN":
                print(f"[draft-section] WARN: {budget_report['warnings']}")
            break

        # Budget FAIL - generate expansion tasks and continue
        if cycle >= max_cycles:
            print(f"[draft-section] max cycles ({max_cycles}) reached, budget still failing")
            break

        # Build expansion tasks from current content
        gap_report = section_gap_report(section, existing_content, chars_per_page, float(args.min_budget_ratio))
        expansion_tasks = gap_report.get("expansion_tasks") or []

        if not expansion_tasks:
            expansion_tasks = build_expansion_tasks(section, drafted_words, chars_per_page)

        if expansion_tasks:
            expansion_history.append({
                "cycle": cycle,
                "drafted_words": drafted_words,
                "expansion_tasks": expansion_tasks,
            })
            print(f"[draft-section] expansion_tasks: {len(expansion_tasks)} items generated")
        else:
            # Nothing specific to expand but still FAIL - stop
            print(f"[draft-section] FAIL but no specific expansion tasks, stopping")
            break

    # Final budget check
    final_words = effective_char_count(strip_markdown_for_count(existing_content))
    final_budget = section_budget_report(
        section, final_words, chars_per_page,
        float(args.min_budget_ratio), float(args.max_budget_ratio)
    )

    # Determine final status
    if final_budget["result"] == "FAIL":
        final_status = "needs_expansion"
        expansion_tasks = build_expansion_tasks(section, final_words, chars_per_page)
        reason = (
            f"达到 max_cycles={max_cycles} 后仍不足；目标字数 {final_budget.get('target_words_min')}，"
            f"90%下限 {final_budget.get('lower_required')}，当前字数 {final_words}，"
            f"缺口 {max((final_budget.get('lower_required') or 0) - final_words, 0)}。"
        )
        append_manual_decision(
            Path(args.manual_decisions),
            section_id,
            str(section.get("title") or ""),
            reason,
            expansion_tasks,
        )
        flags = section_manual_flags(section)
        if "needs_expansion: max_cycles reached below 90% of target_words_min" not in flags:
            flags.append("needs_expansion: max_cycles reached below 90% of target_words_min")
        section["manual_flags"] = flags
        section["notes"] = reason
        print(f"[draft-section] WARNING: final budget FAIL ({final_words} < {final_budget.get('lower_required')}), marking needs_expansion")
    else:
        final_status = "drafted"

    # Update section status
    section["status"] = final_status
    section["completed_at"] = now_iso()
    section["updated_at"] = now_iso()
    section["drafted_words"] = final_words
    section.pop("lock_path", None)

    # Remove lock file
    if lock_path.exists():
        lock_path.unlink()

    write_yaml_like(plan_path, plan)
    write_section_status(plan, status_yml, status_md, locks_dir, sections_dir)

    print(f"[draft-section] completed section {section_id}: {final_status} ({final_words} words)")
    return 0


def _build_draft_section_system_prompt() -> str:
    return """你是一个专业的标书写作助手，负责生成高质量的标书章节正文。

写作要求：
- 只写可见正文，不得生成 <!-- tender:section --> 等元数据注释
- 不得使用四级或更深的标题（#### 及以下）
- 不得照抄招标文件原文，必须转述后点对点响应
- 不得编造未经确认的功能、接口字段、性能指标、人员、业绩
- 不得使用 Web 搜索结果拼凑正文
- 每个结论后面尽量给出支撑：措施、功能、流程、责任人、交付物或验收方式
- 语气应接近有经验标书作者：稳妥、具体、可核查
- 避免"通过……实现……""构建……体系""全面提升……能力"等机械化句式
- 图表位置要有引导语和解释语

输出要求：
- 只输出 Markdown 格式的正文内容
- 不要输出任何解释、说明或元数据
- 直接开始写正文第一个标题
"""


def _build_draft_section_user_prompt(
    section_id: str,
    section_title: str,
    brief_content: str,
    blueprint: list[dict],
    target_min: int,
    chars_per_page: int,
    existing_content: str,
    expansion_history: list[dict],
) -> str:
    parts = [
        f"# 写作任务：章节 {section_id} — {section_title}",
        "",
        "## 章节 Brief（必须先阅读）",
        brief_content,
        "",
    ]

    if existing_content:
        parts.extend([
            "## 现有内容（请在末尾继续扩写，不要重复已有内容）",
            existing_content,
            "",
        ])
        # Show expansion tasks from previous cycles
        for hist in expansion_history:
            tasks = hist.get("expansion_tasks") or []
            if tasks:
                parts.append(f"## 扩写任务（来自上一轮，还需继续完成）")
                for i, task in enumerate(tasks, 1):
                    parts.append(f"{i}. {task}")
                parts.append("")
    else:
        parts.extend([
            "## 写作目标",
            f"本节目标字数：{target_min} 字以上",
            f"（chars_per_page={chars_per_page}，按此换算页数预算）",
            "",
        ])

    if blueprint:
        parts.extend([
            "## 写作蓝图（按顺序展开每个 block）",
            "",
        ])
        for block in blueprint:
            target = block.get("target_words") or ""
            parts.append(f"- **{block['label']}（{block['name']}）**：目标 {target} 字 | {block.get('guidance', '')}")
        parts.append("")

    if expansion_history:
        # Include what was attempted before
        for hist in expansion_history:
            tasks = hist.get("expansion_tasks") or []
            if tasks:
                parts.append(f"**已尝试的扩写任务（供你参考，不要重复相同的动作）**：")
                for task in tasks:
                    parts.append(f"- {task}")
                parts.append("")

    parts.extend([
        "## 开始写作",
        "严格按照上述要求写正文。完成后在末尾加上当前字数统计，格式：",
        f"<!-- words: {existing_content and 'X' or 'Y'} -->",
        "如果目标字数不足，继续扩写直到达标。",
    ])

    return "\n".join(parts)


def _call_llm_generate(system_prompt: str, user_prompt: str, model: str = "") -> str:
    """Call claude -p to generate content."""
    import subprocess
    cmd = ["claude", "-p", "--output-format", "text"]
    if model:
        cmd.extend(["--model", model])
    cmd.append(user_prompt)

    env = {}
    full_system = system_prompt
    proc = subprocess.Popen(
        cmd,
        stdin=subprocess.PIPE,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        env={**__import__("os").environ, **env},
        cwd=None,
    )
    stdout, stderr = proc.communicate(input=full_system.encode("utf-8"))

    if proc.returncode != 0:
        raise RuntimeError(f"claude -p failed: {stderr.decode('utf-8', errors='replace')}")

    return stdout.decode("utf-8", errors="replace").strip()


def cmd_section_gap(args) -> int:
    plan = load_writing_plan(Path(args.writing_plan))
    section = find_section(plan, args.section_id)
    if section is None:
        raise RuntimeError(f"section not found: {args.section_id}")
    section_file = Path(args.section_file) if args.section_file else section_file_path(section, Path(args.sections_dir))
    if not section_file.exists():
        raise FileNotFoundError(f"section file not found: {section_file}")

    report = section_gap_report(
        section,
        read_text(section_file),
        chars_per_page=resolve_chars_per_page_from_args(args),
        min_ratio=float(args.min_budget_ratio),
    )
    write_json(Path(args.out_json), report)
    lines = [
        f"# Section Gap: {report['section_id']} {report['title']}".rstrip(),
        "",
        f"- Result: {report['result']}",
        f"- Drafted words: {report['drafted_words']}",
        f"- Words to lower required: {report['words_to_lower_required']}",
        "",
        "## Missing Blocks",
        "",
    ]
    if report["missing_blocks"]:
        for block in report["missing_blocks"]:
            lines.append(f"- {block['label']} (`{block['name']}`): {block['guidance']}")
    else:
        lines.append("- None.")
    lines.extend(["", "## Expansion Tasks", ""])
    if report["expansion_tasks"]:
        lines.extend([f"- {task}" for task in report["expansion_tasks"]])
    else:
        lines.append("- None.")
    lines.extend(["", "## Writing Blueprint", ""])
    if report["blueprint"]:
        lines.append("| Order | Block | Target words | Guidance |")
        lines.append("|---:|---|---:|---|")
        for block in report["blueprint"]:
            lines.append(
                f"| {block['order']} | {block['label']} (`{block['name']}`) | {block['target_words']} | {block['guidance']} |"
            )
    else:
        lines.append("- None.")
    write_text(Path(args.out_md), "\n".join(lines) + "\n")
    print(f"[OK] wrote: {args.out_json}")
    print(f"[OK] wrote: {args.out_md}")
    if report["result"] == "FAIL" and not args.allow_fail_exit_zero:
        raise RuntimeError("section gap check failed; continue expansion tasks before complete-section or merge-sections")
    return 0


def audit_page_plan(page_plan_path: Path, manifest: dict) -> dict:
    target_pages_min = int(manifest.get("target_pages_min") or 0)
    target_pages_max = int(manifest.get("target_pages_max") or 0)
    target_is_set = bool(target_pages_min or target_pages_max)
    if not page_plan_path.exists():
        result = "NOT_SET" if not target_is_set else "FAIL"
        issues = []
        if target_is_set:
            issues.append("已设置页数要求，但缺少 page_plan.yml；必须先在提纲阶段统筹分配页数。")
        return {
            "path": str(page_plan_path),
            "result": result,
            "status": "MISSING",
            "planned_pages_min": 0,
            "planned_pages_max": 0,
            "allocation_count": 0,
            "issues": issues,
            "items_with_manual_status": [],
        }

    plan = load_yaml_like(page_plan_path)
    if not isinstance(plan, dict):
        return {
            "path": str(page_plan_path),
            "result": "FAIL",
            "status": "INVALID",
            "planned_pages_min": 0,
            "planned_pages_max": 0,
            "allocation_count": 0,
            "issues": ["page_plan.yml 无法解析为字典结构。"],
            "items_with_manual_status": [],
        }

    plan_status = str(plan.get("status") or "PLANNED").strip()
    allocation = plan.get("allocation") or []
    if not isinstance(allocation, list):
        allocation = []

    issues = []
    manual_items = []
    planned_min = 0.0
    planned_max = 0.0
    padding_words = ("凑页", "凑字", "套话", "泛泛", "重复段落", "无来源", "补水")

    if not target_is_set:
        issues.append("未设置 target_pages_min/target_pages_max；页数只能由用户提供，不能自行猜测。")
    if plan_status.upper() in {"NEED_USER_INPUT", "MANUAL", "FAIL"}:
        issues.append(f"page_plan.yml 状态为 {plan_status}，尚未达到可控写作条件。")
    if not allocation:
        issues.append("page_plan.yml 缺少 allocation；必须按章节分配页数。")

    for idx, item in enumerate(allocation, start=1):
        if not isinstance(item, dict):
            issues.append(f"allocation 第 {idx} 项不是字典结构。")
            continue
        section_id = str(item.get("section_id") or "").strip()
        title = str(item.get("title") or "").strip()
        reason = str(item.get("reason") or "").strip()
        status = str(item.get("status") or "OK").strip().upper()
        pages_min = number_value(item.get("pages_min"))
        pages_max = number_value(item.get("pages_max"))
        if status in {"MANUAL", "FAIL", "TBD", "NEED_USER_INPUT"}:
            manual_items.append({"section_id": section_id, "title": title, "status": status})
        if not section_id or not title:
            issues.append(f"allocation 第 {idx} 项缺少 section_id 或 title。")
        if pages_min is None or pages_max is None:
            issues.append(f"allocation 第 {idx} 项缺少 pages_min/pages_max。")
            continue
        if pages_min < 0 or pages_max < 0 or pages_min > pages_max:
            issues.append(f"allocation 第 {idx} 项页数范围无效：{pages_min}~{pages_max}。")
            continue
        planned_min += pages_min
        planned_max += pages_max
        if not reason:
            issues.append(f"allocation 第 {idx} 项缺少分配理由；页数预算必须说明评分、采购需求或交付逻辑依据。")
        if any(word in reason for word in padding_words):
            issues.append(f"allocation 第 {idx} 项疑似以凑页为理由：{reason[:80]}")

    if target_pages_min and planned_min < target_pages_min:
        issues.append(f"页数预算下限 {planned_min:g} 小于用户要求下限 {target_pages_min}。")
    if manual_items:
        issues.append("page_plan.yml 中仍有 MANUAL/FAIL/TBD 项，不能进入 build。")

    result = "PASS"
    if not target_is_set:
        result = "MANUAL"
    if issues:
        result = "FAIL" if target_is_set else "MANUAL"

    return {
        "path": str(page_plan_path),
        "result": result,
        "status": plan_status,
        "planned_pages_min": round(planned_min, 1),
        "planned_pages_max": round(planned_max, 1),
        "allocation_count": len(allocation),
        "issues": issues,
        "items_with_manual_status": manual_items,
    }


def audit_writing_plan_budget(writing_plan_path: Path, manifest: dict, chars_per_page: int) -> dict:
    target_pages_min = int(manifest.get("target_pages_min") or 0)
    target_pages_max = int(manifest.get("target_pages_max") or 0)
    target_is_set = bool(target_pages_min or target_pages_max)
    if not writing_plan_path.exists():
        result = "NOT_SET" if not target_is_set else "FAIL"
        issues = []
        if target_is_set:
            issues.append("已设置页数要求，但缺少 writing_plan.yml；无法验证章节页数/字数预算。")
        return {
            "path": str(writing_plan_path),
            "result": result,
            "section_count": 0,
            "planned_pages_min": 0,
            "planned_pages_max": 0,
            "target_words_min_total": 0,
            "target_words_max_total": 0,
            "issues": issues,
        }

    try:
        plan = load_writing_plan(writing_plan_path)
    except Exception as exc:
        return {
            "path": str(writing_plan_path),
            "result": "FAIL",
            "section_count": 0,
            "planned_pages_min": 0,
            "planned_pages_max": 0,
            "target_words_min_total": 0,
            "target_words_max_total": 0,
            "issues": [f"writing_plan.yml 无法解析：{exc}"],
        }

    issues = []
    planned_min = 0.0
    planned_max = 0.0
    words_min_total = 0
    words_max_total = 0
    section_count = 0

    for idx, section in enumerate(plan.get("sections", []), start=1):
        if not isinstance(section, dict):
            issues.append(f"sections 第 {idx} 项不是字典结构。")
            continue
        section_count += 1
        section_id = section_id_of(section) or f"#{idx}"
        title = str(section.get("title") or "")
        manual_flags = section_manual_flags(section)
        page_min = number_value(section.get("page_budget_min"))
        page_max = number_value(section.get("page_budget_max"))
        explicit_min = number_value(section.get("target_words_min"))
        explicit_max = number_value(section.get("target_words_max"))
        min_words, max_words, _ = section_budget_words(section, chars_per_page)

        if target_is_set and not manual_flags:
            if page_min is None or page_max is None:
                issues.append(f"{section_id} {title}: 缺少 page_budget_min/page_budget_max。")
            elif page_min < 0 or page_max < 0 or page_min > page_max:
                issues.append(f"{section_id} {title}: 页数预算无效：{page_min}~{page_max}。")
            if "req_ids" not in section:
                issues.append(f"{section_id} {title}: 缺少 req_ids。")
            if "scoring_items" not in section:
                issues.append(f"{section_id} {title}: 缺少 scoring_items。")
            if "content_contract" not in section and "required_slots" not in section:
                issues.append(f"{section_id} {title}: 缺少 content_contract 或 required_slots。")

        if page_min is not None:
            planned_min += page_min
        if page_max is not None:
            planned_max += page_max
        if min_words is not None:
            words_min_total += min_words
        if max_words is not None:
            words_max_total += max_words

        if page_min is not None and explicit_min is not None:
            expected_min = int(round(page_min * chars_per_page))
            if int(round(explicit_min)) != expected_min:
                issues.append(
                    f"{section_id} {title}: target_words_min={explicit_min:g} 低于 "
                    f"page_budget_min={page_min:g} 页按 {chars_per_page} 字/页换算的最低预算 {expected_min}；必须精确按页数预算传导。"
                )
        if page_max is not None and explicit_max is not None:
            expected_max = int(round(page_max * chars_per_page))
            if int(round(explicit_max)) != expected_max:
                issues.append(
                    f"{section_id} {title}: target_words_max={explicit_max:g} 不等于 "
                    f"page_budget_max={page_max:g} 页按 {chars_per_page} 字/页换算的预算 {expected_max}；必须精确按页数预算传导。"
                )
        if explicit_min is not None and explicit_max is not None and explicit_min > explicit_max:
            issues.append(f"{section_id} {title}: target_words_min 大于 target_words_max。")

    if target_pages_min and planned_min < target_pages_min:
        issues.append(f"writing_plan 页数预算下限 {planned_min:g} 小于用户要求下限 {target_pages_min}。")
    if target_pages_min and words_min_total < target_pages_min * chars_per_page:
        issues.append(
            f"writing_plan 字数预算下限 {words_min_total:g} 低于用户页数下限换算值 "
            f"{int(target_pages_min * chars_per_page)}。"
        )

    result = "PASS"
    if not target_is_set:
        result = "MANUAL"
        issues.append("未设置 target_pages_min/target_pages_max；预算审计只能检查结构，不能验证总页数。")
    if issues:
        result = "FAIL" if target_is_set else "MANUAL"

    return {
        "path": str(writing_plan_path),
        "result": result,
        "section_count": section_count,
        "planned_pages_min": round(planned_min, 1),
        "planned_pages_max": round(planned_max, 1),
        "target_words_min_total": words_min_total,
        "target_words_max_total": words_max_total,
        "issues": issues,
    }


def cmd_budget_audit(args) -> int:
    manifest = read_manifest(Path(args.manifest))
    chars_per_page = resolve_chars_per_page(args.chars_per_page, manifest)
    page_report = audit_page_plan(Path(args.page_plan), manifest)
    writing_report = audit_writing_plan_budget(Path(args.writing_plan), manifest, chars_per_page)

    result = "PASS"
    blockers = []
    for label, report in (("page_plan", page_report), ("writing_plan", writing_report)):
        if report["result"] in {"FAIL", "MANUAL"}:
            result = "FAIL" if report["result"] == "FAIL" else "MANUAL"
            blockers.append(f"{label} budget check result: {report['result']}")

    report = {
        "result": result,
        "chars_per_page": chars_per_page,
        "target_pages_min": int(manifest.get("target_pages_min") or 0),
        "target_pages_max": int(manifest.get("target_pages_max") or 0),
        "page_plan": page_report,
        "writing_plan": writing_report,
        "blockers": blockers,
    }
    write_json(Path(args.out_json), report)

    md_lines = [
        "# Budget Audit",
        "",
        f"- Result: {result}",
        f"- Chars per page: {chars_per_page}",
        f"- Target pages: {report['target_pages_min'] or 'unset'} ~ {report['target_pages_max'] or 'unset'}",
        f"- Page plan: {page_report['result']} ({page_report['planned_pages_min']} ~ {page_report['planned_pages_max']} pages)",
        f"- Writing plan: {writing_report['result']} ({writing_report['planned_pages_min']} ~ {writing_report['planned_pages_max']} pages, min words {writing_report['target_words_min_total']})",
        "",
        "## Issues",
        "",
    ]
    issues = list(page_report.get("issues") or []) + list(writing_report.get("issues") or [])
    if issues:
        md_lines.extend([f"- {item}" for item in issues])
    else:
        md_lines.append("- None.")
    write_text(Path(args.out_md), "\n".join(md_lines) + "\n")

    print(f"[OK] wrote: {args.out_json}")
    print(f"[OK] wrote: {args.out_md}")
    if result == "FAIL":
        raise RuntimeError("writing plan budget audit failed; fix budget_plan/page_plan/writing_plan before draft")
    return 0


def actual_words_for_section_from_draft(section: dict, draft_text: str) -> int:
    section_id = section_id_of(section)
    title = str(section.get("title") or "").strip()
    headings = list(re.finditer(r"^(#{1,3})\s+(.+?)\s*$", draft_text, flags=re.MULTILINE))
    for idx, heading in enumerate(headings):
        heading_text = heading.group(2).strip()
        if (section_id and section_id in heading_text) or (title and title in heading_text):
            start = heading.start()
            end = headings[idx + 1].start() if idx + 1 < len(headings) else len(draft_text)
            return effective_char_count(strip_markdown_for_count(draft_text[start:end]))
    recorded = number_value(section.get("drafted_words"))
    return int(recorded) if recorded is not None else 0


def build_word_budget_report(
    manifest_path: Path,
    page_plan_path: Path,
    writing_plan_path: Path,
    draft_md_path: Path,
    chars_per_page_arg: object = "",
) -> dict:
    manifest = read_manifest(manifest_path)
    chars_per_page = resolve_chars_per_page(chars_per_page_arg, manifest)
    target_pages_min = int(manifest.get("target_pages_min") or 0)
    target_pages_max = int(manifest.get("target_pages_max") or 0)
    target_words_min = target_pages_min * chars_per_page
    draft_text = read_text(draft_md_path)
    current_words = effective_char_count(strip_markdown_for_count(draft_text))
    page_report = audit_page_plan(page_plan_path, manifest)
    writing_report = audit_writing_plan_budget(writing_plan_path, manifest, chars_per_page)
    try:
        plan = load_writing_plan(writing_plan_path)
    except Exception:
        plan = {"sections": []}

    section_reports = []
    under_budget = []
    for section in plan.get("sections", []):
        if not isinstance(section, dict):
            continue
        section_id = section_id_of(section)
        section_actual = actual_words_for_section_from_draft(section, draft_text)
        min_words, max_words, _ = section_budget_words(section, chars_per_page)
        lower_required = int(math.ceil((min_words or 0) * 0.90)) if min_words else 0
        suggestions = build_expansion_tasks(section, section_actual, chars_per_page)
        item = {
            "section_id": section_id,
            "title": section.get("title", ""),
            "actual_words": section_actual,
            "target_words_min": min_words or 0,
            "target_words_max": max_words or 0,
            "lower_required": lower_required,
            "gap_to_lower_required": max(lower_required - section_actual, 0),
            "gap_to_target_min": max((min_words or 0) - section_actual, 0),
            "suggested_expansion_points": suggestions,
            "result": "FAIL" if lower_required and section_actual < lower_required else "PASS",
        }
        section_reports.append(item)
        if item["result"] == "FAIL":
            under_budget.append(item)

    writing_min_total = int(writing_report.get("target_words_min_total") or 0)
    return {
        "result": "FAIL" if under_budget or current_words < int(target_words_min * 0.90) else "PASS",
        "target_pages_min": target_pages_min,
        "target_pages_max": target_pages_max,
        "chars_per_page": chars_per_page,
        "target_words_min": target_words_min,
        "writing_plan_target_words_min_total": writing_min_total,
        "current_words": current_words,
        "estimated_pages": round(current_words / max(chars_per_page, 1), 1),
        "total_gap_to_target_min": max(target_words_min - current_words, 0),
        "total_gap_to_90_percent": max(int(math.ceil(target_words_min * 0.90)) - current_words, 0),
        "page_plan": page_report,
        "writing_plan": writing_report,
        "sections": section_reports,
        "under_budget_sections": under_budget,
    }


def cmd_word_budget_report(args) -> int:
    report = build_word_budget_report(
        Path(args.manifest),
        Path(args.page_plan),
        Path(args.writing_plan),
        Path(args.draft_md),
        getattr(args, "chars_per_page", ""),
    )
    write_json(Path(args.out_json), report)

    lines = [
        "# Word Budget Report",
        "",
        f"- Result: {report['result']}",
        f"- 目标页数: {report['target_pages_min'] or '未设下限'} ~ {report['target_pages_max'] or '未设上限'}",
        f"- 每页估算字数: {report['chars_per_page']}",
        f"- 目标最低字数: {report['target_words_min']}",
        f"- writing_plan 最低字数合计: {report['writing_plan_target_words_min_total']}",
        f"- 当前已写字数: {report['current_words']}",
        f"- 估算页数: {report['estimated_pages']}",
        f"- 总缺口: {report['total_gap_to_target_min']}",
        "",
        "## 低于预算的章节",
        "",
    ]
    if report["under_budget_sections"]:
        lines.append("| Section | Title | Actual | 90% Required | Target Min | Gap |")
        lines.append("|---|---|---:|---:|---:|---:|")
        for item in report["under_budget_sections"]:
            safe_title = str(item["title"]).replace("|", "\\|")
            lines.append(
                f"| {item['section_id']} | {safe_title} | "
                f"{item['actual_words']} | {item['lower_required']} | {item['target_words_min']} | {item['gap_to_lower_required']} |"
            )
        lines.extend(["", "## 每节建议扩写点", ""])
        for item in report["under_budget_sections"]:
            lines.append(f"### {item['section_id']} {item['title']}".rstrip())
            suggestions = item.get("suggested_expansion_points") or []
            if suggestions:
                lines.extend([f"- {point}" for point in suggestions])
            else:
                lines.append("- 围绕评分项响应、采购需求响应、业务流程、输入输出、接口边界、部署位置、异常处理、交付物和验收依据补足。")
            lines.append("")
    else:
        lines.append("- None.")
    write_text(Path(args.out_md), "\n".join(lines).rstrip() + "\n")
    print(f"[OK] wrote: {args.out_json}")
    print(f"[OK] wrote: {args.out_md}")
    return 0


def cmd_draft_audit(args) -> int:
    draft_path = Path(args.draft_md)
    if not draft_path.exists():
        raise FileNotFoundError(f"Draft Markdown not found: {draft_path}")

    manifest = read_manifest(Path(args.manifest))
    markdown = read_text(draft_path)
    headings = extract_md_headings(markdown)
    deep_headings = [h for h in headings if h["level"] > 3]
    repeated_shell = repeated_document_shell_headings(headings)
    visible_heading_text = "\n".join(h["normalized"] for h in headings if h["level"] <= 3)

    scoring_items = extract_scoring_items(Path(args.scoring_matrix))
    missing_scoring_titles = []
    for item in scoring_items:
        normalized = normalize_title(item)
        if normalized and normalized not in visible_heading_text:
            missing_scoring_titles.append(item)

    plain = strip_markdown_for_count(markdown)
    effective_chars = effective_char_count(plain)
    chars_per_page = resolve_chars_per_page(args.chars_per_page, manifest)
    estimated_pages = round(effective_chars / max(chars_per_page, 1), 1)
    target_pages_min = int(manifest.get("target_pages_min") or 0)
    target_pages_max = int(manifest.get("target_pages_max") or 0)
    max_is_hard = bool(manifest.get("target_pages_max_is_hard") or manifest.get("page_compression_required"))
    page_status = "NOT_SET"
    over_target_pages_max = False
    page_warnings = []
    if target_pages_min or target_pages_max:
        lower_ok = estimated_pages >= target_pages_min if target_pages_min else True
        page_status = "PASS" if lower_ok else "FAIL"
        over_target_pages_max = bool(target_pages_max and estimated_pages > target_pages_max)
        if over_target_pages_max:
            if max_is_hard:
                page_status = "FAIL"
            else:
                if page_status == "PASS":
                    page_status = "WARN"
                page_warnings.append(
                    f"估算页数 {estimated_pages} 超出目标上限 {target_pages_max}，默认仅作为 WARN；"
                    "只有用户明确要求硬上限或压缩页数时才作为阻断。"
                )
    page_plan_report = audit_page_plan(Path(args.page_plan), manifest)
    if hasattr(args, "writing_plan"):
        writing_plan_path = Path(args.writing_plan)
        writing_plan_report = audit_writing_plan_budget(writing_plan_path, manifest, chars_per_page)
        word_budget_report = build_word_budget_report(
            Path(args.manifest),
            Path(args.page_plan),
            writing_plan_path,
            draft_path,
            args.chars_per_page,
        )
    else:
        writing_plan_report = {
            "result": "NOT_SET",
            "target_words_min_total": 0,
            "issues": [],
        }
        word_budget_report = {
            "result": "PASS",
            "current_words": effective_chars,
            "total_gap_to_target_min": 0,
            "under_budget_sections": [],
        }

    result = "PASS"
    blockers = []
    if deep_headings:
        result = "FAIL"
        blockers.append("存在四级或更深标题，评分项可能无法在目录三级内被专家找到。")
    if missing_scoring_titles:
        result = "FAIL"
        blockers.append("存在未进入三级标题内的评分项题目；请回到提纲阶段调整 scoring_title_placement.md、outline.md 和 writing_plan.yml。")
    if repeated_shell:
        result = "FAIL"
        blockers.append("存在重复封面/目录类一级标题，疑似整稿壳或章节重复合并；请先清理 v1.md 或重新运行 merge-sections。")
    if page_status == "FAIL":
        result = "FAIL"
        if over_target_pages_max and max_is_hard:
            blockers.append("估算页数超出用户明确设置的硬性页数上限。")
        else:
            blockers.append("估算页数低于用户提供的页数下限。")
    if page_plan_report["result"] in {"FAIL", "MANUAL"}:
        result = "FAIL" if page_plan_report["result"] == "FAIL" else "MANUAL"
        blockers.append("页数预算未通过统筹检查；请回到 outline 阶段调整 page_plan.yml 和 writing_plan.yml。")
    if hasattr(args, "writing_plan") and writing_plan_report["result"] in {"FAIL", "MANUAL"}:
        result = "FAIL" if writing_plan_report["result"] == "FAIL" else "MANUAL"
        blockers.append("writing_plan.yml 字数预算未通过检查；不得进入 build。")
    if hasattr(args, "writing_plan") and word_budget_report["result"] == "FAIL":
        result = "FAIL"
        blockers.append("v1.md 当前字数低于 page_plan/writing_plan 预算；必须先扩写缺口章节。")

    report = {
        "result": result,
        "draft_md": str(draft_path),
        "headings_total": len(headings),
        "deep_headings": deep_headings,
        "repeated_document_shell": repeated_shell,
        "scoring_items_total": len(scoring_items),
        "missing_scoring_titles": missing_scoring_titles,
        "effective_chars": effective_chars,
        "chars_per_page": chars_per_page,
        "estimated_pages": estimated_pages,
        "target_pages_min": target_pages_min,
        "target_pages_max": target_pages_max,
        "target_pages_max_is_hard": max_is_hard,
        "over_target_pages_max": over_target_pages_max,
        "page_status": page_status,
        "page_warnings": page_warnings,
        "page_plan": page_plan_report,
        "writing_plan": writing_plan_report,
        "word_budget": word_budget_report,
        "blockers": blockers,
        "note": "Page count is an estimate before final Word/PDF layout. Final page numbers must be checked after build/export.",
    }
    write_json(Path(args.out_json), report)

    md_lines = [
        "# 草稿审查摘要",
        "",
        f"- 结果：{result}",
        f"- 标题数量：{len(headings)}",
        f"- 四级及更深标题：{len(deep_headings)}",
        f"- 评分项数量：{len(scoring_items)}",
        f"- 未进入三级标题内的评分项：{len(missing_scoring_titles)}",
        f"- 估算页数：{estimated_pages} 页（按 {chars_per_page} 有效字/页估算）",
        f"- 页数要求：{target_pages_min or '未设下限'} ~ {target_pages_max or '未设上限'} 页",
        f"- 页数上限硬约束：{'是' if max_is_hard else '否'}",
        f"- 页数预算：{page_plan_report['result']}，计划 {page_plan_report['planned_pages_min']} ~ {page_plan_report['planned_pages_max']} 页，章节数 {page_plan_report['allocation_count']}",
        f"- 写作预算：{writing_plan_report['result']}，最低字数合计 {writing_plan_report['target_words_min_total']}",
        f"- 当前字数预算：{word_budget_report['result']}，当前 {word_budget_report['current_words']} 字，总缺口 {word_budget_report['total_gap_to_target_min']} 字",
        "",
        "## 阻断项",
        "",
    ]
    if blockers:
        md_lines.extend([f"- {item}" for item in blockers])
    else:
        md_lines.append("- 暂无。")
    if page_warnings:
        md_lines.extend(["", "## 页数提醒", ""])
        md_lines.extend([f"- {item}" for item in page_warnings])
    if missing_scoring_titles:
        md_lines.extend(["", "## 缺失评分项标题", ""])
        md_lines.extend([f"- {item}" for item in missing_scoring_titles])
        md_lines.extend(
            [
                "",
                "处理要求：如果题目没有合适章节位置，不要硬塞到无关章节；应回到提纲阶段标记 MANUAL，请用户确认是否调整目录、合并到相近题目，或新增三级以内标题。",
            ]
        )
    if deep_headings:
        md_lines.extend(["", "## 过深标题", ""])
        md_lines.extend([f"- L{h['level']} 第 {h['line']} 行：{h['title']}" for h in deep_headings[:80]])
        md_lines.extend(
            [
                "",
                "处理要求：不要一律把四级标题改成加粗正文。先判断该标题是否承载评分题目或专家必须查找的题目；如果承载，必须回到 outline/scoring_title_placement/writing_plan 调整为三级以内标题；只有非评分题目、非目录查找项，才可改为上级标题下的加粗段首、列表项或普通正文。",
            ]
        )
    if repeated_shell:
        md_lines.extend(["", "## 重复封面/目录", ""])
        md_lines.extend(
            [
                f"- {item['title']} 出现 {item['count']} 次，行号："
                + "、".join(str(line) for line in item["lines"])
                for item in repeated_shell
            ]
        )
        md_lines.extend(
            [
                "",
                "处理要求：这通常说明模板骨架、封面目录或旧版 v1.md 被重复并入正文。不要继续 build；先检查 section 文件和 merge-sections 输入，只保留一套封面/目录/正文结构。",
            ]
        )
    if page_plan_report["issues"]:
        md_lines.extend(["", "## 页数预算问题", ""])
        md_lines.extend([f"- {item}" for item in page_plan_report["issues"]])
    if writing_plan_report["issues"]:
        md_lines.extend(["", "## writing_plan 预算问题", ""])
        md_lines.extend([f"- {item}" for item in writing_plan_report["issues"]])
    if word_budget_report["under_budget_sections"]:
        md_lines.extend(["", "## 字数不足章节", ""])
        for item in word_budget_report["under_budget_sections"]:
            md_lines.append(
                f"- {item['section_id']} {item['title']}: 当前 {item['actual_words']}，"
                f"90%下限 {item['lower_required']}，缺口 {item['gap_to_lower_required']}。"
            )
    md_lines.extend(
        [
            "",
            "## 页数处理原则",
            "",
            "- 页数预算必须在提纲阶段完成，先按采购需求、评分标准、主要内容和总页数设计章节篇幅，再进入正文写作。",
            "- 页数不足时，应优先补足已确认的采购需求、评分项、流程、输入输出、交付物和验收依据。",
            "- 页数超出时，默认只检查重复表述、通用套话和与本项目无关内容；只有用户明确要求压缩或设置硬性上限时，才把压缩作为阻断修复。",
            "- 不得为了凑页数添加未确认承诺、无来源材料或与项目无关段落。",
        ]
    )
    write_text(Path(args.out_md), "\n".join(md_lines) + "\n")
    print(f"[OK] wrote: {args.out_json}")
    print(f"[OK] wrote: {args.out_md}")
    return 0


DIAGRAM_REQUIRED_TERMS = (
    "总体架构设计",
    "总体架构",
    "系统架构",
    "技术架构",
    "技术路线",
    "数据流程",
    "数据流",
    "业务流程",
    "模型训练流程",
    "模型训练",
    "部署架构",
    "接口链路",
    "接口设计",
    "系统对接",
    "运维流程",
)


def markdown_has_diagram(text: str) -> bool:
    if "<!-- tender:diagram" in text or "<!-- tender:diagram-source:" in text:
        return True
    if re.search(r"!\[[^\]]*\]\([^)]+\)", text):
        return True
    if re.search(r"<img\b", text, flags=re.IGNORECASE):
        return True
    if re.search(r"```mermaid[\s\S]+?```", text, flags=re.IGNORECASE):
        return True
    return False


def diagram_audit_items(sections_dir: Path, draft_md: Optional[Path] = None) -> list[dict]:
    sources: list[Path] = []
    if sections_dir.exists():
        sources.extend(sorted(path for path in sections_dir.glob("*.md") if path.is_file()))
    if not sources and draft_md and draft_md.exists():
        sources.append(draft_md)

    items = []
    for path in sources:
        text = read_text(path)
        if not text.strip():
            continue
        matched_terms = [term for term in DIAGRAM_REQUIRED_TERMS if term in text]
        if not matched_terms:
            continue
        has_diagram = markdown_has_diagram(text)
        items.append(
            {
                "section_id": path.stem,
                "path": str(path),
                "matched_terms": matched_terms,
                "has_diagram": has_diagram,
                "status": "OK" if has_diagram else "MISSING",
                "recommended_command": f"/tender:diagram-gen {path} --engine fireworks --insert",
            }
        )
    return items


def cmd_diagram_audit(args) -> int:
    sections_dir = Path(args.sections_dir)
    draft_md = Path(args.draft_md) if getattr(args, "draft_md", "") else None
    items = diagram_audit_items(sections_dir, draft_md)
    missing = [item for item in items if item["status"] == "MISSING"]
    report = {
        "result": "WARN" if missing else "PASS",
        "sections_scanned": len(items),
        "missing_count": len(missing),
        "candidates": items,
        "missing_diagrams": missing,
    }
    write_json(Path(args.out_json), report)

    lines = [
        "# 图表审查",
        "",
        f"- 结果：{report['result']}",
        f"- 需要配图章节：{len(items)}",
        f"- 应有图但没有图：{len(missing)}",
        "",
    ]
    if missing:
        lines.extend(["## 应有图但没有图", ""])
        for item in missing:
            terms = "、".join(item["matched_terms"])
            lines.append(f"- {item['section_id']}：{item['path']}；命中：{terms}")
            lines.append(f"  - 建议：{item['recommended_command']}")
        lines.extend(
            [
                "",
                "处理要求：review 必须在生成 review.md 之前补图；如果无法调用 /tender:diagram-gen，必须执行同等逻辑生成 SVG 并插入章节 Markdown。",
            ]
        )
    else:
        lines.append("- 暂无缺图。")
    write_text(Path(args.out_md), "\n".join(lines) + "\n")
    print(f"[OK] wrote: {args.out_json}")
    print(f"[OK] wrote: {args.out_md}")
    return 0


@dataclass
class MdBlock:
    kind: str
    text: str = ""
    level: int = 0
    image_path: str = ""
    caption: str = ""
    rows: Optional[list[list[str]]] = None
    style: Optional[str] = None


def parse_style_comment(comment: str) -> Optional[str]:
    match = re.search(r"style:\s*([^\n\r]+)", comment)
    return match.group(1).strip() if match else None


def split_table_row(line: str) -> list[str]:
    cells = line.strip().strip("|").split("|")
    return [cell.strip() for cell in cells]


def is_table_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.match(r"^:?-{3,}:?$", cell.replace(" ", "")) for cell in cells)


def parse_markdown(markdown: str) -> list[MdBlock]:
    lines = markdown.splitlines()
    blocks: list[MdBlock] = []
    paragraph: list[str] = []
    pending_style: Optional[str] = None
    in_code = False
    code: list[str] = []
    i = 0

    def flush_paragraph() -> None:
        nonlocal paragraph
        text = "\n".join(paragraph).strip()
        if text:
            blocks.append(MdBlock(kind="paragraph", text=text, style=pending_style))
        paragraph = []

    def flush_code() -> None:
        nonlocal code
        if code:
            blocks.append(MdBlock(kind="code", text="\n".join(code), style=pending_style))
        code = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("<!--") and "tender:format" in stripped:
            comment = [stripped]
            while "-->" not in comment[-1] and i + 1 < len(lines):
                i += 1
                comment.append(lines[i])
            pending_style = parse_style_comment("\n".join(comment))
            i += 1
            continue

        if stripped.startswith("```"):
            if in_code:
                in_code = False
                flush_code()
            else:
                flush_paragraph()
                in_code = True
            i += 1
            continue

        if in_code:
            code.append(line)
            i += 1
            continue

        if not stripped:
            flush_paragraph()
            pending_style = None
            i += 1
            continue

        image = re.search(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image:
            flush_paragraph()
            blocks.append(
                MdBlock(kind="image", image_path=image.group(2).strip(), caption=image.group(1).strip(), style=pending_style)
            )
            pending_style = None
            i += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            blocks.append(
                MdBlock(kind="heading", level=len(heading.group(1)), text=heading.group(2).strip(), style=pending_style)
            )
            pending_style = None
            i += 1
            continue

        if re.match(r"^[-*]\s+.+", stripped):
            flush_paragraph()
            blocks.append(MdBlock(kind="bullet", text=stripped[2:].strip(), style=pending_style))
            i += 1
            continue

        if re.match(r"^\d+\.\s+.+", stripped):
            flush_paragraph()
            blocks.append(MdBlock(kind="number", text=re.sub(r"^\d+\.\s+", "", stripped), style=pending_style))
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            flush_paragraph()
            rows = [split_table_row(stripped)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_table_row(lines[i]))
                i += 1
            blocks.append(MdBlock(kind="table", rows=rows, style=pending_style))
            pending_style = None
            continue

        paragraph.append(line)
        i += 1

    flush_paragraph()
    if in_code:
        flush_code()
    return blocks


def find_placeholder(doc: Document, placeholder: str) -> Optional[Paragraph]:
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            return paragraph
    return None


def insert_paragraph_before(paragraph: Paragraph, text: str = "", style: Optional[str] = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    if text:
        new_para.add_run(text)
    return new_para


def set_cell_text(cell, text: str) -> None:
    cell.text = text


def insert_table_before(paragraph: Paragraph, rows: list[list[str]], style: Optional[str]) -> Table:
    column_count = max((len(row) for row in rows), default=1)
    table = paragraph._parent.add_table(rows=len(rows), cols=column_count, width=Inches(6.2))
    if style:
        try:
            table.style = style
        except Exception:
            pass
    for r_idx, row in enumerate(rows):
        for c_idx in range(column_count):
            set_cell_text(table.cell(r_idx, c_idx), row[c_idx] if c_idx < len(row) else "")
    paragraph._p.addprevious(table._tbl)
    return table


def remove_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)
    paragraph._p = paragraph._element = None


def style_for_block(block: MdBlock, style_map: dict) -> Optional[str]:
    if block.style:
        return block.style
    if block.kind == "heading":
        if block.level <= 1:
            return style_map.get("h1", DEFAULT_STYLE_MAP["h1"])
        if block.level == 2:
            return style_map.get("h2", DEFAULT_STYLE_MAP["h2"])
        if block.level == 3:
            return style_map.get("h3", DEFAULT_STYLE_MAP["h3"])
        return style_map.get("h4", style_map.get("h3", DEFAULT_STYLE_MAP["h3"]))
    return style_map.get(block.kind, style_map.get("paragraph", DEFAULT_STYLE_MAP["paragraph"]))


def validate_style_map(doc: Document, style_map: dict) -> None:
    available = {style_name(style): style for style in doc.styles}
    required_roles = {
        "h1": WD_STYLE_TYPE.PARAGRAPH,
        "h2": WD_STYLE_TYPE.PARAGRAPH,
        "h3": WD_STYLE_TYPE.PARAGRAPH,
        "paragraph": WD_STYLE_TYPE.PARAGRAPH,
        "bullet": WD_STYLE_TYPE.PARAGRAPH,
        "number": WD_STYLE_TYPE.PARAGRAPH,
        "caption": WD_STYLE_TYPE.PARAGRAPH,
        "table": WD_STYLE_TYPE.TABLE,
    }
    missing = []
    wrong_type = []
    for role, expected_type in required_roles.items():
        mapped = style_map.get(role)
        if not mapped:
            missing.append(f"{role}: <empty>")
            continue
        style = available.get(mapped)
        if style is None:
            missing.append(f"{role}: {mapped}")
            continue
        try:
            if expected_type is not None and style.type != expected_type:
                wrong_type.append(f"{role}: {mapped}")
        except Exception:
            pass
    if missing or wrong_type:
        details = []
        if missing:
            details.append("missing styles: " + "; ".join(missing))
        if wrong_type:
            details.append("style type mismatch: " + "; ".join(wrong_type))
        raise RuntimeError(
            "Template style map is not build-ready. Run /tender:template with a complete Word template "
            "so styles are auto-detected from the selected template. " + " | ".join(details)
        )


def resolve_image_path(image_path: str, draft_path: Path, figures_base: Path) -> Path:
    path = Path(image_path.strip().strip('"').strip("'"))
    if path.is_absolute():
        return path
    candidates = [
        figures_base / path,
        draft_path.parent / path,
        Path.cwd() / path,
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return candidates[0]


def build_docx_from_markdown(
    manifest_path: Path,
    draft_md: Path,
    out_docx: Path,
    format_rules_path: Optional[Path],
    template_override: Optional[Path],
) -> dict:
    require_docx()
    manifest = read_manifest(manifest_path)
    template_path = template_override or Path(manifest.get("template_docx_path", "inputs/template.docx"))
    placeholder = manifest.get("body_placeholder", "{{TECH_SECTION_BODY}}")
    figures_base = Path(manifest.get("figures_out_path", "work/50_figures/out"))

    if not template_path.exists():
        raise FileNotFoundError(f"Template DOCX not found: {template_path}")
    if not draft_md.exists():
        raise FileNotFoundError(f"Draft Markdown not found: {draft_md}")

    rules = {}
    rules_source = ""
    if format_rules_path and format_rules_path.exists():
        rules = read_json(format_rules_path, {})
        rules_source = str(format_rules_path)
    elif manifest.get("template_profile_path") and Path(manifest["template_profile_path"]).exists():
        rules = read_json(Path(manifest["template_profile_path"]), {})
        rules_source = str(manifest["template_profile_path"])
    else:
        raise RuntimeError(
            "Template format rules are missing. Run /tender:template before /tender:build so fonts and styles come from the selected template."
        )

    style_map = deepcopy(DEFAULT_STYLE_MAP)
    if isinstance(rules, dict):
        if "style_map" in rules:
            style_map.update(rules.get("style_map") or {})
        if "format_rules" in rules and isinstance(rules["format_rules"], dict):
            style_map.update((rules["format_rules"].get("style_map") or {}))

    doc = Document(str(template_path))
    placeholder_paragraph = find_placeholder(doc, placeholder)
    if placeholder_paragraph is None:
        raise RuntimeError(f"Placeholder not found in template DOCX: {placeholder}")
    validate_style_map(doc, style_map)

    markdown = draft_md.read_text(encoding="utf-8")
    blocks = parse_markdown(markdown)
    figure_no = 1
    table_no = 1

    for block in blocks:
        style = style_for_block(block, style_map)
        if block.kind in {"heading", "paragraph", "bullet", "number", "code"}:
            insert_paragraph_before(placeholder_paragraph, block.text, style)
        elif block.kind == "image":
            image_file = resolve_image_path(block.image_path, draft_md, figures_base)
            p = insert_paragraph_before(placeholder_paragraph, "", None)
            if image_file.exists():
                try:
                    p.add_run().add_picture(str(image_file), width=Inches(6.2))
                except Exception:
                    p.add_run().add_picture(str(image_file))
            else:
                p.add_run(f"[图像不存在：{block.image_path}]")
            caption = block.caption or image_file.stem
            insert_paragraph_before(placeholder_paragraph, f"图{figure_no}  {caption}", style_map.get("caption"))
            figure_no += 1
        elif block.kind == "table" and block.rows:
            insert_table_before(placeholder_paragraph, block.rows, style_map.get("table"))
            caption = block.caption or f"表{table_no}"
            if caption:
                insert_paragraph_before(placeholder_paragraph, caption, style_map.get("caption"))
            table_no += 1

    remove_paragraph(placeholder_paragraph)
    ensure_parent(out_docx)
    doc.save(str(out_docx))

    return {
        "template_docx": str(template_path),
        "draft_md": str(draft_md),
        "out_docx": str(out_docx),
        "placeholder": placeholder,
        "format_rules_source": rules_source,
        "style_map": style_map,
        "blocks": len(blocks),
        "figures": figure_no - 1,
        "tables": table_no - 1,
        "manual_after_build": [
            "Open Word and update the table of contents.",
            "Refresh page numbers in scoring and mandatory-response indexes.",
            "Manually inspect landscape tables and section breaks if wide tables were used.",
        ],
    }


def cmd_build_docx(args) -> int:
    manifest = read_manifest(Path(args.manifest))
    draft_path = Path(args.draft_md)
    markdown = read_text(draft_path)
    repeated_shell = repeated_document_shell_headings(extract_md_headings(markdown))
    if repeated_shell:
        details = "; ".join(
            f"{item['title']} x{item['count']} at lines {','.join(str(line) for line in item['lines'])}"
            for item in repeated_shell
        )
        raise RuntimeError(
            "draft structure check failed before DOCX build: repeated cover/TOC shell detected. "
            "Clean v1.md or rerun merge-sections before build. "
            + details
        )
    effective_chars = effective_char_count(strip_markdown_for_count(markdown))
    chars_per_page = resolve_chars_per_page(args.chars_per_page, manifest)
    estimated_pages = round(effective_chars / max(chars_per_page, 1), 1)
    target_pages_min = int(manifest.get("target_pages_min") or 0)
    target_pages_max = int(manifest.get("target_pages_max") or 0)
    if target_pages_min or target_pages_max:
        lower_ok = estimated_pages >= target_pages_min if target_pages_min else True
        if not lower_ok:
            raise RuntimeError(
                "draft page budget check failed before DOCX build: "
                f"estimated_pages={estimated_pages}, "
                f"target_pages={target_pages_min or 'unset'}~{target_pages_max or 'unset'}, "
                f"effective_chars={effective_chars}, chars_per_page={chars_per_page}. "
                "Run review/draft-audit, then revise under the relevant scoring items and procurement "
                "requirements, or mark unsupported gaps as MANUAL instead of building a short final document."
            )

    report = build_docx_from_markdown(
        manifest_path=Path(args.manifest),
        draft_md=Path(args.draft_md),
        out_docx=Path(args.out_docx),
        format_rules_path=Path(args.format_rules) if args.format_rules else None,
        template_override=Path(args.template) if args.template else None,
    )
    if args.out_log:
        write_text(
            Path(args.out_log),
            "# DOCX 构建日志\n\n"
            + "\n".join([f"- {key}: {value}" for key, value in report.items() if key != "manual_after_build"])
            + "\n\n## 生成后人工确认\n\n"
            + "\n".join([f"- {item}" for item in report["manual_after_build"]])
            + "\n",
        )
    print(f"[OK] wrote: {args.out_docx}")
    return 0


def cmd_progress_seed(args) -> int:
    tasks_path = Path(args.tasks)
    progress_path = Path(args.progress)
    if not tasks_path.exists():
        tasks = [
            {
                "id": f"{idx:02d}_{name}",
                "stage": name,
                "title": title,
                "status": "pending",
                "updated_at": "",
                "notes": "",
            }
            for idx, (name, title) in enumerate(WORKFLOW_STAGES, start=1)
        ]
        if yaml is not None:
            tasks_text = yaml.safe_dump({"tasks": tasks}, allow_unicode=True, sort_keys=False)
        else:
            tasks_text = json.dumps({"tasks": tasks}, ensure_ascii=False, indent=2)
        write_text(tasks_path, tasks_text)
    if not progress_path.exists():
        lines = [
            "# Tender 进度记录",
            "",
            "| 阶段 | 状态 | 产物 | 备注 |",
            "|---|---|---|---|",
        ]
        for name, title in WORKFLOW_STAGES:
            lines.append(f"| {name} | pending |  | {title} |")
        write_text(progress_path, "\n".join(lines) + "\n")
    print(f"[OK] ensured: {tasks_path}")
    print(f"[OK] ensured: {progress_path}")
    return 0


RESET_STAGE_OUTPUTS = {
    "template": ["20_templates", "20_requirements/format_rules.json"],
    "extract": ["10_source_extracted", "20_requirements"],
    "brief": ["15_user_brief"],
    "outline": ["30_plan"],
    "draft": ["40_drafts", "50_figures"],
    "review": ["70_review"],
    "build": ["60_build"],
}


def reset_stage_names_from(to_stage: str) -> list[str]:
    ordered = [name for name, _ in WORKFLOW_STAGES if name != "init"]
    if to_stage not in ordered:
        raise RuntimeError(f"unknown reset stage: {to_stage}. Expected one of: {', '.join(ordered)}")
    return ordered[ordered.index(to_stage):]


def stage_index(stage: str) -> int:
    ordered = [name for name, _ in WORKFLOW_STAGES]
    if stage not in ordered:
        raise RuntimeError(f"unknown stage: {stage}")
    return ordered.index(stage)


def unique_archive_dir(base: Path) -> Path:
    if not base.exists():
        return base
    suffix = 2
    while True:
        candidate = Path(f"{base}-{suffix}")
        if not candidate.exists():
            return candidate
        suffix += 1


def compact_relative_paths(relative_paths: list[str]) -> list[str]:
    paths = [Path(item) for item in relative_paths]
    compacted = []
    for idx, path in enumerate(paths):
        is_child_of_other = False
        for other_idx, other in enumerate(paths):
            if idx == other_idx:
                continue
            if len(other.parts) < len(path.parts) and path.parts[: len(other.parts)] == other.parts:
                is_child_of_other = True
                break
        if not is_child_of_other:
            compacted.append(path.as_posix())
    return compacted


def archive_relative_paths(work_dir: Path, archive_dir: Path, relative_paths: Iterable[str]) -> list[str]:
    archived = []
    for relative in relative_paths:
        source = work_dir / relative
        if not source.exists():
            continue
        destination = archive_dir / relative
        ensure_parent(destination)
        shutil.move(str(source), str(destination))
        archived.append(relative)
    return archived


def update_tasks_after_reset(tasks_path: Path, reset_from_stage: str, timestamp: str, archive_dir: Path) -> None:
    tasks_doc = load_yaml_like(tasks_path)
    if not isinstance(tasks_doc, dict):
        return
    tasks = tasks_doc.get("tasks")
    if not isinstance(tasks, list):
        return
    reset_stages = set(reset_stage_names_from(reset_from_stage))
    for task in tasks:
        if not isinstance(task, dict):
            continue
        stage = str(task.get("stage") or "")
        if stage in reset_stages:
            task["status"] = "pending"
            task["updated_at"] = timestamp
            task["notes"] = f"Reset to {reset_from_stage}; previous outputs archived at {archive_dir}"
    write_yaml_like(tasks_path, tasks_doc)


def write_progress_after_reset(progress_path: Path, to_stage: str, timestamp: str, archive_dir: Path, archived: list[str]) -> None:
    lines = [
        "# Tender 进度记录",
        "",
        f"- 最近回退：{timestamp} reset to `{to_stage}`",
        f"- 归档目录：`{archive_dir}`",
        f"- 归档产物：{len(archived)} 项",
        "",
        "| 阶段 | 状态 | 产物 | 备注 |",
        "|---|---|---|---|",
    ]
    reset_stages = set(reset_stage_names_from(to_stage))
    for name, title in WORKFLOW_STAGES:
        status = "pending" if name in reset_stages else "kept"
        lines.append(f"| {name} | {status} |  | {title} |")
    write_text(progress_path, "\n".join(lines) + "\n")


def write_context_handoff_after_reset(path: Path, to_stage: str, timestamp: str, archive_dir: Path, archived: list[str]) -> None:
    lines = [
        "# Tender Context Handoff",
        "",
        f"- Updated at: {timestamp}",
        f"- Current state: reset to `{to_stage}`",
        f"- Archived previous outputs: `{archive_dir}`",
        "",
        "## Next Step",
        "",
        f"- Rerun `/tender:{to_stage}` before using later-stage outputs.",
        "- In a new session, read this file, `work/00_manifest.yml`, and the current stage inputs before continuing.",
        "",
        "## Archived Paths",
        "",
    ]
    if archived:
        lines.extend([f"- `{item}`" for item in archived])
    else:
        lines.append("- None.")
    write_text(path, "\n".join(lines) + "\n")


def cmd_reset_stage(args) -> int:
    work_dir = Path(args.work_dir)
    to_stage = str(args.to_stage or "").strip()
    timestamp = str(args.timestamp or "").strip() or datetime.now().strftime("%Y%m%d-%H%M%S")
    archive_root = Path(args.archive_root) if args.archive_root else work_dir / "90_archive"
    archive_dir = unique_archive_dir(archive_root / f"{timestamp}-reset-to-{to_stage}")

    reset_stages = reset_stage_names_from(to_stage)
    relative_paths = []
    seen = set()
    for stage in reset_stages:
        for relative in RESET_STAGE_OUTPUTS.get(stage, []):
            if relative not in seen:
                relative_paths.append(relative)
                seen.add(relative)
    relative_paths = compact_relative_paths(relative_paths)

    archived = archive_relative_paths(work_dir, archive_dir, relative_paths)

    manifest_path = Path(args.manifest)
    manifest = read_manifest(manifest_path)
    if stage_index(to_stage) <= stage_index("outline"):
        manifest["planning_completed"] = False
    if stage_index(to_stage) <= stage_index("draft"):
        manifest["drafting_completed"] = False
        manifest["draft_cursor"] = ""
        manifest["drafted_words"] = 0
    manifest["last_reset"] = {
        "to_stage": to_stage,
        "reset_at": timestamp,
        "archive_dir": str(archive_dir),
        "archived_paths": archived,
    }
    write_manifest(manifest_path, manifest)

    update_tasks_after_reset(Path(args.tasks), to_stage, timestamp, archive_dir)
    write_progress_after_reset(Path(args.progress), to_stage, timestamp, archive_dir, archived)
    write_context_handoff_after_reset(Path(args.context_handoff), to_stage, timestamp, archive_dir, archived)

    print(f"[OK] reset to stage: {to_stage}")
    print(f"[OK] archived {len(archived)} paths to: {archive_dir}")
    print(f"[OK] wrote: {args.context_handoff}")
    return 0


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(prog="tenderctl")
    sub = parser.add_subparsers(dest="cmd", required=True)

    sp = sub.add_parser("template-profile", help="Inspect DOCX template and export JSON + Markdown")
    sp.add_argument("--template", required=True)
    sp.add_argument("--template-id", default="")
    sp.add_argument("--out-json", default="work/20_templates/default/template_profile.json")
    sp.add_argument("--out-md", default="work/20_templates/default/template.md")
    sp.add_argument("--out-format-rules", default="")
    sp.add_argument("--manifest", default="work/00_manifest.yml")
    sp.add_argument("--build-template-out", default="")
    sp.add_argument("--insert-after", default="")
    sp.add_argument("--build-mode", choices=["local_template", "replace_body", "insert_after"], default="")
    sp.add_argument("--template-scope", choices=["local", "complete"], default="")
    sp.add_argument("--body-start-after", default="")
    sp.add_argument("--body-end-before", default="")
    sp.set_defaults(func=cmd_template_profile)

    sp = sub.add_parser("ingest", help="Extract text from one PDF/DOCX input")
    sp.add_argument("--manifest", default="work/00_manifest.yml")
    sp.add_argument("--tender", default="")
    sp.add_argument("--out-txt", default="work/10_source_extracted/tender.txt")
    sp.set_defaults(func=cmd_ingest)

    sp = sub.add_parser("ingest-inputs", help="Extract tender text and optional separate scoring-standard text")
    sp.add_argument("--manifest", default="work/00_manifest.yml")
    sp.add_argument("--out-tender-txt", default="work/10_source_extracted/tender.txt")
    sp.add_argument("--out-scoring-txt", default="work/10_source_extracted/scoring_standard.txt")
    sp.add_argument("--out-user-brief-txt", default="work/10_source_extracted/user_brief.txt")
    sp.set_defaults(func=cmd_ingest_inputs)

    sp = sub.add_parser("brief-check", help="Check user brief relevance and deviation risks against tender text")
    sp.add_argument("--tender-txt", default="work/10_source_extracted/tender.txt")
    sp.add_argument("--user-brief-txt", default="work/10_source_extracted/user_brief.txt")
    sp.add_argument("--out-json", default="work/15_user_brief/brief_alignment.json")
    sp.add_argument("--out-md", default="work/15_user_brief/brief_alignment.md")
    sp.add_argument("--warn-threshold", default="0.12")
    sp.set_defaults(func=cmd_brief_check)

    sp = sub.add_parser("reset-stage", help="Archive current and later stage outputs so a tender stage can be rerun cleanly")
    sp.add_argument("--work-dir", default="work")
    sp.add_argument("--to-stage", choices=[name for name, _ in WORKFLOW_STAGES if name != "init"], required=True)
    sp.add_argument("--archive-root", default="")
    sp.add_argument("--manifest", default="work/00_manifest.yml")
    sp.add_argument("--tasks", default="work/00_tasks.yml")
    sp.add_argument("--progress", default="work/00_progress.md")
    sp.add_argument("--context-handoff", default="work/00_context_handoff.md")
    sp.add_argument("--timestamp", default="")
    sp.set_defaults(func=cmd_reset_stage)

    sp = sub.add_parser("claim-section", help="Atomically claim a draft section for parallel writing")
    sp.add_argument("--writing-plan", default="work/30_plan/writing_plan.yml")
    sp.add_argument("--section-status", default="work/40_drafts/section_status.yml")
    sp.add_argument("--out-md", default="work/40_drafts/section_status.md")
    sp.add_argument("--locks-dir", default="work/40_drafts/locks")
    sp.add_argument("--sections-dir", default="work/40_drafts/v1_sections")
    sp.add_argument("--section-briefs-dir", default="work/40_drafts/section_briefs")
    sp.add_argument("--section-id", default="")
    sp.add_argument("--owner", default="agent")
    sp.add_argument("--allow-existing-file", action="store_true")
    sp.add_argument("--skip-section-brief", action="store_true")
    sp.set_defaults(func=cmd_claim_section)

    sp = sub.add_parser("complete-section", help="Mark a claimed draft section as drafted or blocked")
    sp.add_argument("--manifest", default="work/00_manifest.yml")
    sp.add_argument("--writing-plan", default="work/30_plan/writing_plan.yml")
    sp.add_argument("--section-status", default="work/40_drafts/section_status.yml")
    sp.add_argument("--out-md", default="work/40_drafts/section_status.md")
    sp.add_argument("--locks-dir", default="work/40_drafts/locks")
    sp.add_argument("--sections-dir", default="work/40_drafts/v1_sections")
    sp.add_argument("--section-id", required=True)
    sp.add_argument("--section-file", default="")
    sp.add_argument("--owner", default="")
    sp.add_argument("--status", choices=["drafted", "blocked", "MANUAL", "needs_revision", "needs_expansion"], default="drafted")
    sp.add_argument("--words", default="")
    sp.add_argument("--chars-per-page", default="")
    sp.add_argument("--min-budget-ratio", default="0.90")
    sp.add_argument("--max-budget-ratio", default="1.10")
    sp.add_argument("--notes", default="")
    sp.add_argument("--manual-decisions", default="work/30_plan/manual_decisions.md")
    sp.add_argument("--allow-no-lock", action="store_true")
    sp.add_argument("--keep-lock", action="store_true")
    sp.set_defaults(func=cmd_complete_section)

    sp = sub.add_parser("merge-sections", help="Merge drafted section files into the main draft")
    sp.add_argument("--manifest", default="work/00_manifest.yml")
    sp.add_argument("--writing-plan", default="work/30_plan/writing_plan.yml")
    sp.add_argument("--section-status", default="work/40_drafts/section_status.yml")
    sp.add_argument("--out-md", default="work/40_drafts/section_status.md")
    sp.add_argument("--sections-dir", default="work/40_drafts/v1_sections")
    sp.add_argument("--locks-dir", default="work/40_drafts/locks")
    sp.add_argument("--out-draft", default="work/40_drafts/v1.md")
    sp.add_argument("--allow-partial", action="store_true")
    sp.add_argument("--chars-per-page", default="")
    sp.add_argument("--min-budget-ratio", default="0.90")
    sp.add_argument("--max-budget-ratio", default="1.10")
    sp.set_defaults(func=cmd_merge_sections)

    sp = sub.add_parser("normalize-writing-plan", help="Normalize writing_plan.yml sections to list form")
    sp.add_argument("--writing-plan", default="work/30_plan/writing_plan.yml")
    sp.set_defaults(func=cmd_normalize_writing_plan)

    sp = sub.add_parser("content-contract", help="Add hidden writing contracts to writing_plan.yml without changing headings")
    sp.add_argument("--manifest", default="work/00_manifest.yml")
    sp.add_argument("--writing-plan", default="work/30_plan/writing_plan.yml")
    sp.add_argument("--page-plan", default="work/30_plan/page_plan.yml")
    sp.add_argument("--scoring-matrix", default="work/20_requirements/scoring_matrix.md")
    sp.add_argument("--mandatory-matrix", default="work/20_requirements/mandatory_matrix.md")
    sp.add_argument("--out-writing-plan", default="")
    sp.add_argument("--chars-per-page", default="")
    sp.set_defaults(func=cmd_content_contract)

    sp = sub.add_parser("section-gap", help="Audit one drafted section and generate expansion tasks")
    sp.add_argument("--manifest", default="work/00_manifest.yml")
    sp.add_argument("--writing-plan", default="work/30_plan/writing_plan.yml")
    sp.add_argument("--sections-dir", default="work/40_drafts/v1_sections")
    sp.add_argument("--section-id", required=True)
    sp.add_argument("--section-file", default="")
    sp.add_argument("--out-json", default="work/70_review/section_gap.json")
    sp.add_argument("--out-md", default="work/70_review/section_gap.md")
    sp.add_argument("--chars-per-page", default="")
    sp.add_argument("--min-budget-ratio", default="0.90")
    sp.add_argument("--allow-fail-exit-zero", action="store_true")
    sp.set_defaults(func=cmd_section_gap)

    sp = sub.add_parser("draft-section", help="End-to-end section drafting with automatic budget enforcement loop")
    sp.add_argument("--manifest", default="work/00_manifest.yml")
    sp.add_argument("--writing-plan", default="work/30_plan/writing_plan.yml")
    sp.add_argument("--section-status", default="work/40_drafts/section_status.yml")
    sp.add_argument("--out-md", default="work/40_drafts/section_status.md")
    sp.add_argument("--locks-dir", default="work/40_drafts/locks")
    sp.add_argument("--sections-dir", default="work/40_drafts/v1_sections")
    sp.add_argument("--section-briefs-dir", default="work/40_drafts/section_briefs")
    sp.add_argument("--section-id", default="")
    sp.add_argument("--owner", default="agent")
    sp.add_argument("--model", default="")
    sp.add_argument("--allow-existing-file", action="store_true")
    sp.add_argument("--continue-existing", action="store_true")
    sp.add_argument("--chars-per-page", default="")
    sp.add_argument("--min-budget-ratio", default="0.90")
    sp.add_argument("--max-budget-ratio", default="1.10")
    sp.add_argument("--max-cycles", default="3")
    sp.add_argument("--manual-decisions", default="work/30_plan/manual_decisions.md")
    sp.set_defaults(func=cmd_draft_section)

    sp = sub.add_parser("similarity", help="Check local source/draft similarity")
    sp.add_argument("--tender-txt", default="work/10_source_extracted/tender.txt")
    sp.add_argument("--draft-md", default="work/40_drafts/v1.md")
    sp.add_argument("--out-json", default="work/40_drafts/v1.similarity.json")
    sp.add_argument("--k", default="10")
    sp.add_argument("--window", default="400")
    sp.add_argument("--step", default="80")
    sp.add_argument("--threshold", default="0.15")
    sp.set_defaults(func=cmd_similarity)

    sp = sub.add_parser("budget-audit", help="Audit page_plan and writing_plan page/word budget alignment")
    sp.add_argument("--manifest", default="work/00_manifest.yml")
    sp.add_argument("--page-plan", default="work/30_plan/page_plan.yml")
    sp.add_argument("--writing-plan", default="work/30_plan/writing_plan.yml")
    sp.add_argument("--out-json", default="work/30_plan/budget_audit.json")
    sp.add_argument("--out-md", default="work/30_plan/budget_audit.md")
    sp.add_argument("--chars-per-page", default="")
    sp.set_defaults(func=cmd_budget_audit)

    sp = sub.add_parser("word-budget-report", help="Report total and per-section word budget gaps")
    sp.add_argument("--manifest", default="work/00_manifest.yml")
    sp.add_argument("--page-plan", default="work/30_plan/page_plan.yml")
    sp.add_argument("--writing-plan", default="work/30_plan/writing_plan.yml")
    sp.add_argument("--draft-md", default="work/40_drafts/v1.md")
    sp.add_argument("--out-md", default="work/70_review/word_budget_report.md")
    sp.add_argument("--out-json", default="work/70_review/word_budget_report.json")
    sp.add_argument("--chars-per-page", default="")
    sp.set_defaults(func=cmd_word_budget_report)

    sp = sub.add_parser("draft-audit", help="Audit draft headings, scoring-title visibility, and page budget")
    sp.add_argument("--manifest", default="work/00_manifest.yml")
    sp.add_argument("--draft-md", default="work/40_drafts/v1.md")
    sp.add_argument("--scoring-matrix", default="work/20_requirements/scoring_matrix.md")
    sp.add_argument("--page-plan", default="work/30_plan/page_plan.yml")
    sp.add_argument("--writing-plan", default="work/30_plan/writing_plan.yml")
    sp.add_argument("--out-json", default="work/70_review/draft_audit.json")
    sp.add_argument("--out-md", default="work/70_review/draft_audit.md")
    sp.add_argument("--chars-per-page", default="")
    sp.set_defaults(func=cmd_draft_audit)

    sp = sub.add_parser("diagram-audit", help="Audit sections that likely need diagrams before review")
    sp.add_argument("--sections-dir", default="work/40_drafts/v1_sections")
    sp.add_argument("--draft-md", default="work/40_drafts/v1.md")
    sp.add_argument("--out-json", default="work/70_review/diagram_audit.json")
    sp.add_argument("--out-md", default="work/70_review/diagram_audit.md")
    sp.set_defaults(func=cmd_diagram_audit)

    sp = sub.add_parser("build-docx", help="Build final DOCX from Markdown and template")
    sp.add_argument("--manifest", default="work/00_manifest.yml")
    sp.add_argument("--draft-md", default="work/40_drafts/v1.md")
    sp.add_argument("--out-docx", default="work/60_build/final.docx")
    sp.add_argument("--format-rules", default="work/20_requirements/format_rules.json")
    sp.add_argument("--template", default="")
    sp.add_argument("--out-log", default="work/60_build/build_log.md")
    sp.add_argument("--chars-per-page", default="")
    sp.set_defaults(func=cmd_build_docx)

    sp = sub.add_parser("progress-seed", help="Create resumable task/progress files")
    sp.add_argument("--tasks", default="work/00_tasks.yml")
    sp.add_argument("--progress", default="work/00_progress.md")
    sp.set_defaults(func=cmd_progress_seed)

    return parser


def main(argv: Optional[list[str]] = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)
    return args.func(args)


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
